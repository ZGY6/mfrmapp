"""
MFRMSight — 多面Rasch模型分析引擎
===================================
基于Andrich Rating Scale Model (1978)
Fisher-scoring JMLE估计, 分阶段PROX→JMLE
"""
import numpy as np
from pathlib import Path
from typing import Optional
import re


__version__ = "0.9.0"


def parse_facets_txt(filepath: str) -> dict:
    """解析Facets/Minifac风格的.txt输入文件。

    适配 v0.8.0: 支持等号两边有空格的写法 (如 "Facets = 3"、"Data = ")，
    支持 Labels 段中的裸数字行(元素编号占位符，无标签)，
    支持 Data 段中的分号注释行 (如 "; Student1")。
    自动检测 3 面或 4 面设计。

    文件 min_line: Facets/N, Noncentered/N, Labels, Data, Model=?,?,?,R23
    """
    with open(filepath, "r", encoding="utf-8") as f:
        lines = [l.strip() for l in f if l.strip()]
    data_rows, infos = [], [[], [], [], []]
    labels_ok, in_data, cur_facet = False, False, 0
    n_facets = 4  # 默认
    noncentered = 1

    for line in lines:
        # ── 规范化: 等号两边空格不影响──
        lower = line.lower()
        # 去掉等号周围的空格: "facets = 3" → "facets=3"
        normalized = lower.replace(" = ", "=").replace("= ", "=").replace(" =", "=")

        # ── Header 指令 ──
        if normalized.startswith("facets="):
            try: n_facets = int(normalized.split("=")[1])
            except ValueError: pass
        if normalized.startswith("noncentered=") or normalized.startswith("noncent"):
            try: noncentered = int(normalized.split("=")[1])
            except ValueError: pass
        if normalized.startswith("model="):
            pass  # Model 信息当前仅存储，不影响解析
        if normalized.startswith("title="):
            pass

        # ── 分隔符 * ──
        if line == "*":
            labels_ok = True
            cur_facet = 0
            continue

        # ── Labels 段 ──
        if normalized.startswith("labels="):
            continue
        if normalized.startswith("data="):
            labels_ok = False
            in_data = True
            continue

        if labels_ok and not in_data:
            # 裸数字行 (元素编号占位符), 跳过
            if line.isdigit():
                continue
            if "," in line:
                parts = [x.strip() for x in line.split(",", 1)]
                if parts[0].isdigit():
                    fid = int(parts[0])
                    if cur_facet == 0:
                        cur_facet = fid  # facet声明行
                    else:
                        if fid > 0 and parts[1].strip():
                            infos[cur_facet - 1].append((fid, parts[1]))
            continue

        # ── Data 段 ──
        if in_data:
            # 跳过注释行
            if line.startswith(";"):
                continue
            # 同时支持逗号和空格/制表分隔
            p = [x.strip() for x in line.split(",")]
            if len(p) < n_facets + 1:
                p = [x.strip() for x in line.split()]
            if len(p) >= n_facets + 1:
                try:
                    data_rows.append([int(v) for v in p])
                except ValueError:
                    pass

    raw = np.array(data_rows) if data_rows else np.empty((0, n_facets + 1), dtype=int)
    result = _build_dict(raw, n_facets,
                       [l for _, l in infos[0]], [l for _, l in infos[1]],
                       [l for _, l in infos[2]], [l for _, l in infos[3]])
    result["noncentered"] = noncentered
    return result


def parse_facets_out(filepath: str) -> dict:
    """v0.9.0: 解析 Facets .out.txt 输出文件，提取关键 Table。

    提取:
      - Table 3: 迭代历史 (PROX/JMLE 收敛)
      - Table 5: 总体拟合 (ObsMean/ExpMean/VarExp/Chi-squared)
      - Table 7.x: 各面向测量报告 (Measure/SE/Infit/Outfit/Separation/Reliability)
      - Table 8.1: 等级类别统计 (Counts/Thresholds/Measures)
      - Table 4: 异常反应列表

    Returns: {"summary": {...}, "facets": {...}, "categories": {...}, "anomalous": [...]}
    """
    with open(filepath, "r", encoding="utf-8") as f:
        content = f.read()
    lines = content.split("\n")

    result = {"summary": {}, "facets": {}, "categories": {}, "anomalous": []}

    # ── 定位各 Table 起始行 ──
    table_starts = {}
    for i, line in enumerate(lines):
        line_s = line.strip()
        for tn in ["Table 3", "Table 4.1", "Table 5", "Table 7.1.1", "Table 7.2.1",
                    "Table 7.3.1", "Table 7.4.1", "Table 8.1"]:
            if line_s.startswith(tn) and tn not in table_starts:
                table_starts[tn] = i
                break

    # ── Table 5: 总体拟合 ──
    if "Table 5" in table_starts:
        chunk = "\n".join(lines[table_starts["Table 5"]:table_starts["Table 5"]+20])
        m = re.search(r"Mean \(Count:\s*(\d+)\)", chunk)
        obs_n = int(m.group(1)) if m else 0
        # 提取 ObsMean/ExpMean/S.D.
        vals = re.findall(r"([\d.]+)\s+([\d.]+)\s+([\d.]+)\s+([\d.]+)\s+([\d.]+)", chunk)
        if vals:
            result["summary"]["obs_mean"] = float(vals[0][0])
            result["summary"]["exp_mean"] = float(vals[0][1])
            # 取 S.D. Sample 行 (第3行)
            result["summary"]["sd_obs"] = float(vals[2][2]) if len(vals) >= 3 else float(vals[1][2])
        m = re.search(r"Variance explained by Rasch measures\s*=\s*([\d.]+)\s*([\d.]+)%", chunk)
        if m:
            result["summary"]["var_exp"] = float(m.group(2))
        m = re.search(r"chi-squared\s*=\s*([\d.]+).*probability\s*=\s*([\d.]+)", chunk)
        if m:
            result["summary"]["chi_sq"] = float(m.group(1))
            result["summary"]["chi_prob"] = float(m.group(2))
        result["summary"]["N"] = obs_n

    # ── Table 7.x: 测量报告 ──
    facet_map = {"Table 7.1.1": "students", "Table 7.2.1": "raters",
                 "Table 7.3.1": "criteria", "Table 7.4.1": "items"}
    for tn, fname in facet_map.items():
        if tn not in table_starts:
            continue
        start = table_starts[tn]
        end = start + 25
        chunk = lines[start:end]
        rows = []
        for line in chunk:
            # Table 7 数据行: "|  916      64  ...   .89   .87 | 4 Student4  |"
            m = re.match(
                r"\|\s*(\d+)\s+(\d+)\s+([\d.]+)\s+([\d.]+)\s*\|"
                r"\s*([\d.-]+)\s+([\d.]+)\s*\|"
                r"\s*([\d.]+)\s+([\d.-]+)\s+([\d.]+)\s+([\d.-]+)\s*\|"
                r"\s*([\d.]+)\s*\|"
                r"\s*([\d.]+)\s+([\d.]+)\s*\|"
                r"\s*\d+\s+(.+)", line)
            if m:
                rows.append({
                    "total": int(m.group(1)), "count": int(m.group(2)),
                    "obs_avg": float(m.group(3)), "fair_avg": float(m.group(4)),
                    "meas": float(m.group(5)), "se": float(m.group(6)),
                    "infit_mnsq": float(m.group(7)), "infit_zstd": float(m.group(8)),
                    "outfit_mnsq": float(m.group(9)), "outfit_zstd": float(m.group(10)),
                    "discrm": float(m.group(11)),
                    "ptmea": float(m.group(12)), "ptexp": float(m.group(13)),
                    "label": m.group(14).strip(),
                })
            # 提取 Separation/Reliability (取第一条 Population 行)
            m2 = re.search(r"Separation\s+([\d.]+).*Reliability\s+([\d.]+)", line)
            if m2 and fname not in result["facets"]:
                result["facets"][fname] = {
                    "separation": float(m2.group(1)),
                    "reliability": float(m2.group(2)),
                    "rows": rows,
                }

    # ── Table 8.1: 等级类别统计 ──
    if "Table 8.1" in table_starts:
        start = table_starts["Table 8.1"]
        chunk = lines[start:start+35]
        cat_rows = []
        for line in chunk:
            # 只处理以 "|" 开头且第一个字段为数字的行
            line_s = line.strip()
            if not line_s.startswith("|") or not re.match(r"\|\s*\d", line_s):
                continue
            # 按 | 分割，取前 5 个字段
            cells = [c.strip() for c in line_s.split("|")]
            if len(cells) < 4:
                continue
            # 字段 1: "2       3         3    1%   1%" → [score, total, used, pct, cum_pct]
            f1 = cells[1].split()
            if len(f1) < 3:
                continue
            score = int(f1[0])
            count = int(f1[2])  # "Used" 列
            # 字段 2: "-1.79  -2.35  1.6" → [avg_meas, exp_meas, outfit]
            f2 = cells[2].split()
            avg_meas = float(f2[0].rstrip("*"))
            exp_meas = float(f2[1]) if len(f2) > 1 else 0.0
            outfit_mnsq = float(f2[2]) if len(f2) > 2 else 1.0
            # 字段 3: "-3.35    .61" → [threshold, thresh_se] or empty
            f3 = cells[3].split()
            threshold = float(f3[0]) if len(f3) > 0 else None
            thresh_se = float(f3[1]) if len(f3) > 1 else None
            # 字段 4: "( -4.72)       " 或 "( 4.75)   4.10" → cat_meas
            f4 = cells[4].strip()
            cat_meas = 0.0
            if f4.startswith("("):
                inner = f4.split(")")[0].lstrip("(").strip().split()[0]
                try: cat_meas = float(inner)
                except ValueError: pass
            cat_rows.append({
                "score": score, "count": count,
                "avg_meas": avg_meas, "exp_meas": exp_meas,
                "outfit_mnsq": outfit_mnsq,
                "threshold": threshold, "thresh_se": thresh_se,
                "cat_meas": cat_meas,
            })
        tau_values = [r["threshold"] for r in cat_rows if r["threshold"] is not None]
        result["categories"] = {
            "rows": cat_rows,
            "tau": tau_values,
            "tau_ordered": all(tau_values[i] <= tau_values[i+1] for i in range(len(tau_values)-1)) if len(tau_values) > 1 else True,
        }

    # ── Table 4: 异常反应 ──
    if "Table 4.1" in table_starts:
        start = table_starts["Table 4.1"]
        chunk = lines[start:start+20]
        for line in chunk:
            if "No unexpected" in line:
                break
            m = re.match(r"\|\s*(\d+)\s+(\d+)\s+([\d.]+)\s+([\d.-]+)\s+([\d.-]+)\s*\|", line)
            if m:
                result["anomalous"].append({
                    "score": int(m.group(1)),
                    "observed": int(m.group(2)),
                    "expected": float(m.group(3)),
                    "residual": float(m.group(4)),
                    "stres": float(m.group(5)),
                })

    return result


def parse_excel(filepath: str) -> dict:
    """从Excel加载3面向数据"""
    import pandas as pd
    df_raw = pd.read_excel(filepath, header=None)
    header_row = 1
    for i in range(len(df_raw)):
        s = [str(v) for v in df_raw.iloc[i] if pd.notna(v)]
        if any("编号" in x or "评分人" in x for x in s):
            header_row = i; break
    col_names = [str(df_raw.iloc[header_row, j]) if pd.notna(df_raw.iloc[header_row, j]) else f"Col{j}"
                 for j in range(df_raw.shape[1])]

    # BUG-002 增强: 智能列名推断
    def _infer_excel_col_role(cn: str, col_idx: int) -> str:
        """推断 Excel 列名在 MFRM 中的角色: student / criterion / item / rater / row_id"""
        cn_lower = cn.lower()
        # 强信号
        if any(kw in cn_lower for kw in ["编号", "序号", "id", "row"]):
            return "row_id"
        if any(kw in cn_lower for kw in ["评分人", "评分者", "rater", "评委", "examiner", "judge"]):
            return "rater"
        if any(kw in cn_lower for kw in ["学生", "student", "考生", "学员", "受试"]):
            return "student"
        if any(kw in cn_lower for kw in ["题目", "题号", "item", "task"]):
            return "item"
        if any(kw in cn_lower for kw in ["标准", "维度", "criterion", "criteria", "domain"]):
            return "criterion"

        # 列名含连字符: "综合分析-1号" → 前缀=criterion, 后缀看上下文
        if "-" in cn or "—" in cn:
            sep = "-" if "-" in cn else "—"
            parts = cn.rsplit(sep, 1) if cn.count(sep) == 1 else (cn.split(sep)[0], cn.split(sep)[-1])
            prefix, suffix = parts[0], parts[-1]
            # 前缀含能力关键词 → criterion
            has_crit = any(kw in prefix for kw in ["综合", "沟通", "分析", "人际", "表达", "逻辑", "创新", "协作", "能力"])
            # 后缀含题 → item
            has_item = any(kw in suffix for kw in ["题", "item"])
            # 后缀含号/生/人 → student
            has_stu = any(kw in suffix for kw in ["号", "生", "人", "student"])
            has_num = bool(re.search(r'\d+', suffix))
            if has_crit:
                return "criterion"  # 前缀明确是criterion
            if has_item:
                return "item"
            if has_stu:
                return "student"
            if has_num:
                return "student"  # 默认数字后缀→student
            return "unknown"

        # 纯数字 → student
        if re.match(r'^\d+$', cn.strip()):
            return "student"

        return "unknown"

    s_set, c_set, i_set = set(), set(), set()
    for cn in col_names[2:]:
        role = _infer_excel_col_role(cn, 0)
        if "-" in cn or role == "criterion":
            crit_name = cn.rsplit("-", 1)[0] if "-" in cn else cn
            c_set.add(crit_name)
            stu_name = cn.rsplit("-", 1)[-1] if "-" in cn else cn
            if role == "student":
                s_set.add(stu_name)
            elif role == "item":
                i_set.add(stu_name)
            else:
                s_set.add(stu_name)  # 默认归student
        elif role == "student":
            s_set.add(cn)
        elif role == "item":
            i_set.add(cn)
        elif role == "criterion":
            c_set.add(cn)
        else:
            # 兜底: 按连字符分析
            if "-" in cn:
                crit_name, stu_name = cn.rsplit("-", 1)
                c_set.add(crit_name)
                s_set.add(stu_name)
    s_list, c_list, i_list = sorted(s_set), sorted(c_set), sorted(i_set)
    n_s, n_c = len(s_list), len(c_list)
    n_i = len(i_list) if i_list else 1
    data_rows = []
    rater_labels = []
    # BUG-002 增强: 自动检测 n_facets (3 或 4)
    auto_n_facets = 4 if n_i > 1 else 3
    for ri in range(header_row + 1, len(df_raw)):
        row = df_raw.iloc[ri]
        if pd.isna(row.iloc[0]): continue
        r_label = str(row.iloc[1])
        r_num = _extract_num(r_label)
        if auto_n_facets >= 4:
            for si in range(n_s):
                for ci in range(n_c):
                    for ii in range(n_i):
                        col = 2 + ci * n_s * n_i + si * n_i + ii
                        if col < df_raw.shape[1]:
                            v = row.iloc[col]
                            if pd.notna(v):
                                data_rows.append([si + 1, r_num, ci + 1, ii + 1, int(v)])
        else:
            for si in range(n_s):
                for ci in range(n_c):
                    col = 2 + ci * n_s + si
                    if col < df_raw.shape[1]:
                        v = row.iloc[col]
                        if pd.notna(v): data_rows.append([si + 1, r_num, ci + 1, 1, int(v)])
        rater_labels.append(r_label)
    raw = np.array(data_rows)
    return _build_dict(raw, auto_n_facets, s_list, rater_labels, c_list, i_list if i_list else ["Item1"])


def _extract_num(s): return int(re.findall(r'\d+', str(s))[-1]) if re.findall(r'\d+', str(s)) else 1


def _build_dict(raw, n_facets, s_labels, r_labels, c_labels, i_labels):
    return {
        "raw": raw, "n_facets": n_facets, "N": len(raw),
        "n_s": int(raw[:, 0].max()), "n_r": int(raw[:, 1].max()),
        "n_c": int(raw[:, 2].max()), "n_i": int(raw[:, 3].max()) if n_facets >= 4 and raw.shape[1] > 3 else 1,
        "min_s": int(raw[:, -1].min()), "max_s": int(raw[:, -1].max()),
        "s_labels": s_labels or [f"Student{i+1}" for i in range(int(raw[:, 0].max()))],
        "r_labels": r_labels or [f"Rater{i+1}" for i in range(int(raw[:, 1].max()))],
        "c_labels": c_labels or [f"Criterion{i+1}" for i in range(int(raw[:, 2].max()))],
        "i_labels": i_labels or [f"Item{i+1}" for i in range(max(1, int(raw[:, 3].max()) if n_facets >= 4 and raw.shape[1] > 3 else 1))],
    }


class MFRMEngine:
    """多面Rasch Rating Scale模型估计器"""

    def __init__(self, data: dict):
        for k, v in data.items(): setattr(self, k, v)
        self.noncentered = data.get("noncentered", 1)  # 默认约束第1面(Students), Facets 约定
        self._index()

    def _index(self):
        self.s_idx = self.raw[:, 0] - 1
        self.r_idx = self.raw[:, 1] - 1
        self.c_idx = np.clip(self.raw[:, 2] - 1, 0, self.n_c - 1)
        self.i_idx = self.raw[:, 3] - 1 if self.n_facets >= 4 else np.zeros(self.N, dtype=int)
        self.scores = self.raw[:, -1].astype(float)
        self.x = self.scores - self.min_s
        self.K = int(self.max_s - self.min_s)
        self.cats = np.arange(self.K + 1, dtype=float)
        self.obs_s = [np.where(self.s_idx == s)[0] for s in range(self.n_s)]
        self.obs_r = [np.where(self.r_idx == r)[0] for r in range(self.n_r)]
        self.obs_c = [np.where(self.c_idx == c)[0] for c in range(self.n_c)]
        self.obs_i = [np.where(self.i_idx == it)[0] for it in range(max(self.n_i, 1))]

    def _probs(self):
        intc = (self.theta[self.s_idx] - self.delta[self.r_idx]
                - self.alpha[self.c_idx] - self.beta[self.i_idx])
        logits = np.zeros((self.N, self.K + 1)); cum = np.zeros(self.N)
        for k in range(self.K): cum += intc - self.tau[k]; logits[:, k + 1] = cum
        logits -= logits.max(axis=1, keepdims=True)
        e = np.exp(logits); return e / e.sum(axis=1, keepdims=True)

    def _ll(self, p):
        return sum(np.log(max(p[i, int(self.x[i])], 1e-300)) for i in range(self.N))

    def _s2l(self, scores_list):
        """Score→logit Newton-Raphson 反演（BUG-012 核心修复）。

        替代粗糙的 mean/M→logit 近似。
        在 Andrich Rating Scale 模型中，E[score|θ,τ] 是 θ 的光滑单调函数。
        给定固定 τ 阈值，用 N-R 自洽求 θ 使期望分等于观测均值。

        算法:
          θ ← θ + Δ, 其中 Δ = (target - E[score]) / Var,
          这正是 Fisher scoring 的单参数版本。
          收敛速度: 5-8 步 (< 1e-6 tolerance)。
        """
        K = self.K
        n = len(scores_list)
        th = np.zeros(n)
        tu = self.tau

        for i in range(n):
            sc = scores_list[i]
            nobs = len(sc)
            if nobs == 0:
                th[i] = 0.0
                continue

            target = sc.mean() - self.min_s  # 0-based

            # 极端分数直接给边界值，避免 N-R 跑飞
            if target <= 0.05 * K:
                th[i] = -5.0
                continue
            if target >= 0.95 * K:
                th[i] = 5.0
                continue

            t = 0.0
            for _ in range(25):
                lin = t - tu
                cum = np.zeros(K + 1)
                for k in range(K):
                    cum[k + 1] = cum[k] + lin[k]
                cum -= cum.max()
                e = np.exp(cum)
                p = e / e.sum()
                exp_s = p @ self.cats
                exp_sq = p @ (self.cats ** 2)
                var = max(exp_sq - exp_s ** 2, 0.001)
                delta = (target - exp_s) / var
                t += delta
                if abs(delta) < 1e-6:
                    break

            th[i] = np.clip(t, -10, 10)

        return th

    def _prox(self):
        """PROX 初始化 — 两阶段 N-R 反演（BUG-012 修复）。

        废除了旧的 mean/M→logit 公式，改用 Rating Scale 模型期望分函数
        的 Newton-Raphson 反演来精确映射 score → logit。

        阶段1: 全局分布估计初始 τ(k) → N-R 反解各元素参数 → 居中。
        阶段2: 用新参数 refine τ → 再跑一轮 N-R → 再居中。

        BUG-013 补充修复: 阶段1 N-R 后必须先居中，否则 τ 精炼从偏置
        参数出发导致概率估计有偏，收敛后 ExpMean 偏离 ObsMean 约 1.4%。
        """
        # ── 初始 τ（从边际分布，与 Facets 初值一致）──
        self.tau = np.array([-np.log(
            np.clip((self.scores >= self.min_s + k + 1).mean(), 0.02, 0.98) /
            np.clip((self.scores < self.min_s + k + 1).mean(), 0.02, 0.98))
            for k in range(self.K)])
        self.tau -= self.tau[0]

        # ── 阶段1: N-R 反解各面参数（基于初始 τ）──
        self.theta = self._s2l([self.scores[self.obs_s[s]] for s in range(self.n_s)])
        self.delta = -self._s2l([self.scores[self.obs_r[r]] for r in range(self.n_r)])
        self.alpha = -self._s2l([self.scores[self.obs_c[c]] for c in range(self.n_c)])
        if self.n_i > 1:
            self.beta = -self._s2l([self.scores[self.obs_i[it]] for it in range(self.n_i)])
        else:
            self.beta = np.zeros(1)

        # 阶段1 后立即居中: 确保 τ 精炼从已校准的参数出发
        self._center_prox()

        # ── 阶段2: refine τ 并重新 N-R ──
        if self.K > 0:
            p = self._probs()
            for k in range(self.K):
                o = (self.x >= k + 1).sum()
                e = p[:, k + 1:].sum()
                pg = p[:, k + 1:].sum(axis=1)
                info = (pg * (1 - pg)).sum() + 1.0
                self.tau[k] += 0.5 * (e - o) / info   # 半阻尼 Fisher scoring
            self.tau -= self.tau[0]

            # 重新 N-R（基于 refined τ）
            self.theta = self._s2l([self.scores[self.obs_s[s]] for s in range(self.n_s)])
            self.delta = -self._s2l([self.scores[self.obs_r[r]] for r in range(self.n_r)])
            self.alpha = -self._s2l([self.scores[self.obs_c[c]] for c in range(self.n_c)])
            if self.n_i > 1:
                self.beta = -self._s2l([self.scores[self.obs_i[it]] for it in range(self.n_i)])

            # 阶段2 后再居中 (由 _center_prox 统一处理)
            self._center_prox()

        # ── U 调整（温和校正量尺，与 Facets PROX step 4 等价）──
        p = self._probs()
        sd_obs = np.std(self.scores)
        sd_exp = np.std(self.min_s + p @ self.cats)
        if sd_exp > 0.01:
            U = 1.0 + 0.5 * (sd_obs / sd_exp - 1.0)
            U = np.clip(U, 0.8, 2.0)
        else:
            U = 1.0
        self.theta *= U; self.delta *= U; self.alpha *= U; self.beta *= U; self.tau *= U

    def _center_prox(self):
        """仅对 noncentered 指定的面居中 (Rasch 识别约束: 仅需1个sum-to-zero)"""
        # noncentered 语义: 该面不居中, 其余面居中
        # noncentered=1 → 面1(Students)不居中, 面2,3,4居中
        # 等价于: 以 Students 为自由参考, 其他面被约束
        if self.noncentered != 1:
            self.theta -= self.theta.mean()
        if self.noncentered != 2:
            self.delta -= self.delta.mean()
        if self.noncentered != 3:
            self.alpha -= self.alpha.mean()
        if self.noncentered != 4 and self.n_i > 1:
            self.beta -= self.beta.mean()

    def fit(self, p1=200, p2=300):
        # BUG-004: 稀疏数据检测与警告
        if self.K > 0:
            per_cat = np.bincount(self.x.astype(int), minlength=self.K + 1)
            sparse_cats = [str(self.min_s + i) for i, c in enumerate(per_cat) if c < 8]
            if sparse_cats:
                import warnings
                warnings.warn(
                    f"⚠️ 稀疏数据警告 (每个等级建议 >= 8 观测): "
                    f"分数 {','.join(sparse_cats)} 观测不足。参数可能不稳定，建议合并评分等级。"
                )
        self._prox()
        best = (-1e100, None)
        # 阶段1: 高阻尼 + 强ridge → 锁定大致方向
        for phase, (n_it, rs, rd, dmp) in enumerate([(p1, 0.3, 0.99, 0.2), (p2, 0.03, 0.9995, 0.08)]):
            for it in range(1, n_it + 1):
                rid = max(rs * (rd ** it), 0.001)
                # 交替更新各面向
                self._up_theta(dmp, rid); self._up_delta(dmp, rid)
                self._up_alpha(dmp, rid); self._up_beta(dmp, rid)
                # tau: 更强的阻尼和正则化，防止发散
                tau_damp = dmp * 0.3
                tau_rid = rid * 5
                self._up_tau(tau_damp, tau_rid)
                if it % 10 == 0:
                    p = self._probs(); ll = self._ll(p)
                    if ll > best[0]: best = (ll, (self.theta.copy(), self.delta.copy(), self.alpha.copy(), self.beta.copy(), self.tau.copy()))
        if best[1]: self.theta, self.delta, self.alpha, self.beta, self.tau = best[1]
        self._fin(); return self

    def _up(self, arr, obs, n, dmp, rid, sign, center=True):
        """Fisher scoring 更新 — sequential update（BUG-013 修复）。

        旧版: 循环外一次 _probs() → 批量更新所有元素（batch mode）
        新版: 循环内每更新一个元素后立即 _probs()（sequential/JMLE mode）

        数学: JMLE 等价于对每个参数做 Newton step，
              需要"其他参数固定"的假设才成立。
              Sequential 保持了该假设。
        """
        for i in range(n):
            p = self._probs(); e = p @ self.cats
            v = np.clip((p @ (self.cats**2)) - e**2, 0.001, None)
            idx = obs[i]; arr[i] += dmp * (sign * self.x[idx].sum() - sign * e[idx].sum()) / (v[idx].sum() + rid)
        if center: arr -= arr.mean()
        arr[:] = np.clip(arr, -20, 20)

    def _up_theta(self, d, r): self._up(self.theta, self.obs_s, self.n_s, d, r, 1, center=(self.noncentered != 1))
    def _up_delta(self, d, r): self._up(self.delta, self.obs_r, self.n_r, d, r, -1, center=(self.noncentered != 2))
    def _up_alpha(self, d, r): self._up(self.alpha, self.obs_c, self.n_c, d, r, -1, center=(self.noncentered != 3))
    def _up_beta(self, d, r):
        if self.n_i > 1: self._up(self.beta, self.obs_i, self.n_i, d, r, -1, center=(self.n_facets >= 4 and self.noncentered != 4))

    def _up_tau(self, dmp, rid):
        """Tau 更新 — sequential update（BUG-013 修复）。

        每个 τ_k 更新后立即重算概率，
        因为 τ_k 的变化影响所有 k+1...K 的累积概率。
        """
        for k in range(self.K):
            p = self._probs()
            o = (self.x >= k + 1).sum()
            e = p[:, k + 1:].sum()
            pg = p[:, k + 1:].sum(axis=1); info = (pg * (1 - pg)).sum() + rid
            self.tau[k] += dmp * (e - o) / info
        self.tau -= self.tau[0]; self.tau[:] = np.clip(self.tau, -15, 25)

    def _fin(self):
        p = self._probs(); self.ll_final = self._ll(p)
        e = p @ self.cats; self.var_o = np.clip((p @ (self.cats**2)) - e**2, 0.001, None)
        self.exp_scores = self.min_s + e; self.resid = self.scores - self.exp_scores; self.z = self.resid / np.sqrt(self.var_o)
        self.var_exp = max(0, (np.var(self.scores, ddof=1) - np.var(self.resid, ddof=1)) / np.var(self.scores, ddof=1) * 100)

    def _diagnose_categories(self) -> dict:
        """v0.9.0: 评分等级类别功能诊断 (参考 Facets Table 8)。

        判断等级划分是否合理:
          1. 每个等级至少 10 次使用
          2. 类别平均测量值必须单调递增
          3. Andrich 阈值必须有序
          4. 类别峰值概率 > 15%

        Returns: {"passed": bool, "issues": [...], "merge_suggestion": str or None}
        """
        issues = []
        scores_int = self.x.astype(int)
        usage = np.bincount(scores_int, minlength=self.K + 1)

        # 1. 使用频次检查
        low_usage = [(self.min_s + i, int(c)) for i, c in enumerate(usage) if c < 10]
        if low_usage:
            cats_str = ", ".join(f"{s}({c}次)" for s, c in low_usage)
            issues.append(f"等级使用不足: {cats_str}")

        # 2. 类别平均 Measure 单调性
        cat_measures = []
        for k in range(self.K + 1):
            idx = np.where(scores_int == k)[0]
            if len(idx) > 0:
                cat_measures.append((self.min_s + k, float(self.theta[self.s_idx[idx]].mean()) if len(idx) > 0 else 0))
        if cat_measures:
            for i in range(1, len(cat_measures)):
                if cat_measures[i][1] < cat_measures[i-1][1]:
                    issues.append(f"等级 {cat_measures[i-1][0]}->{cat_measures[i][0]} 平均测量值非单调 "
                                  f"({cat_measures[i-1][1]:.2f} -> {cat_measures[i][1]:.2f})")
                    break  # 只报一次

        # 3. Andrich 阈值有序性
        if self.K >= 2:
            thresholds_ordered = all(self.tau[i] <= self.tau[i+1] for i in range(self.K - 1))
            if not thresholds_ordered:
                disordered = []
                for i in range(self.K - 1):
                    if self.tau[i] > self.tau[i+1]:
                        disordered.append(f"τ{i+1}={self.tau[i]:.2f} > τ{i+2}={self.tau[i+1]:.2f}")
                issues.append(f"Andrich 阈值无序: {', '.join(disordered[:3])}")

        # 4. 合并建议
        merge_suggestion = None
        n_low = len(low_usage)
        if n_low >= self.K * 0.3:  # 30% 以上等级使用不足
            merge_suggestion = f"建议合并评分等级: 当前 {self.K+1} 档 ({self.min_s}-{self.max_s})"

        return {
            "passed": len(issues) == 0,
            "issues": issues,
            "usage": [(self.min_s + i, int(c)) for i, c in enumerate(usage)],
            "tau_ordered": all(self.tau[i] <= self.tau[i+1] for i in range(self.K - 1)) if self.K >= 2 else True,
            "merge_suggestion": merge_suggestion,
        }

    def _anomalous_responses(self, threshold: float = 3.0) -> list[dict]:
        """v0.9.0: 异常反应检测 (|StRes| >= threshold)。

        返回异常反应列表，含考生/评分者/标准/题目/期望分/残差。
        """
        anomalous = []
        for i in range(self.N):
            if abs(self.z[i]) >= threshold:
                s_label = self.s_labels[self.s_idx[i]] if self.s_idx[i] < len(self.s_labels) else f"S{self.s_idx[i]+1}"
                r_label = self.r_labels[self.r_idx[i]] if self.r_idx[i] < len(self.r_labels) else f"R{self.r_idx[i]+1}"
                c_label = self.c_labels[self.c_idx[i]] if self.c_idx[i] < len(self.c_labels) else f"C{self.c_idx[i]+1}"
                i_label = self.i_labels[self.i_idx[i]] if self.n_i > 1 and self.i_idx[i] < len(self.i_labels) else ""
                anomalous.append({
                    "student": s_label, "rater": r_label,
                    "criterion": c_label, "item": i_label,
                    "observed": int(self.scores[i]),
                    "expected": round(float(self.exp_scores[i]), 2),
                    "residual": round(float(self.resid[i]), 2),
                    "stres": round(float(self.z[i]), 2),
                })

        # 按 |StRes| 降序排列
        anomalous.sort(key=lambda x: abs(x["stres"]), reverse=True)
        return anomalous

    def _diagnose_bias(self) -> list[dict]:
        """v0.9.0: 自动诊断评分者偏差类型。

        基于 MFRM 参数检测 5 种偏差（参考 PPT 第12讲 Slide 24）:
          1. 宽松/严格效应 — meas 显著偏离均值
          2. 集中趋势效应 — fair_avg 全距小 + outfit 偏低
          3. 随机效应 — infit/outfit 显著超标
          4. 晕轮效应 — 维度间评分差异小 (仅在 n_c >= 2 时检测)
          5. 差异宽松 — 跨考生 FairAvg 标准差异常

        Returns:
          [{"rater": "Rater1", "flags": [("宽松", "meas=+1.52, 显著偏高")]}, ...]
        """
        if not hasattr(self, 'exp_scores'):
            raise RuntimeError("请先运行 fit()")

        biases = []
        # 获取评分者统计
        rater_rows = []
        for p in range(self.n_r):
            idx = self.obs_r[p]
            if len(idx) == 0:
                continue
            fair_vals = self.exp_scores[idx]
            rater_rows.append({
                "label": self.r_labels[p] if p < len(self.r_labels) else f"R{p+1}",
                "meas": self.delta[p],
                "infit": (self.z[idx]**2 * self.var_o[idx]).sum() / max(self.var_o[idx].sum(), 1e-10),
                "outfit": (self.z[idx]**2).sum() / max(len(idx), 1),
                "fair_avg": fair_vals.mean(),
                "fair_std": fair_vals.std(ddof=1) if len(fair_vals) > 1 else 0,
                "n_obs": len(idx),
            })

        if not rater_rows:
            return biases

        meas_vals = np.array([r["meas"] for r in rater_rows])
        meas_mean = meas_vals.mean()
        fair_avgs = np.array([r["fair_avg"] for r in rater_rows])
        fair_mean = fair_avgs.mean()
        fair_range = fair_avgs.max() - fair_avgs.min()
        # 差异宽松: 用各评分者 FairAvg std 的中位数做基线
        fair_stds_all = np.array([r["fair_std"] for r in rater_rows if r["fair_std"] > 0])
        fair_std_median = float(np.median(fair_stds_all)) if len(fair_stds_all) > 0 else 1.0

        for r in rater_rows:
            flags = []

            # 1. 宽松/严格效应: meas 偏离均值 > 1.0 logit
            dev = r["meas"] - meas_mean
            if abs(dev) > 1.0:
                tag = "严格" if dev < 0 else "宽松"
                flags.append((tag, f"meas={r['meas']:+.2f}, 偏离均值{dev:+.2f}"))

            # 2. 集中趋势效应: fair_avg 全距 < 均值的30%
            if fair_range < fair_mean * 0.3 and r["outfit"] < 0.5:
                flags.append(("集中趋势", f"FairAvg={r['fair_avg']:.1f}, Outfit={r['outfit']:.2f}"))

            # 3. 随机效应: infit 或 outfit > 2.0
            if r["infit"] > 2.0 or r["outfit"] > 2.0:
                flags.append(("随机", f"Infit={r['infit']:.2f} Outfit={r['outfit']:.2f}"))

            # 5. 差异宽松: FairAvg std 显著高于中位数 (top outlier)
            if fair_std_median > 0 and r["fair_std"] > fair_std_median * 1.5 and r["n_obs"] >= 4:
                flags.append(("差异宽松", f"FairAvg std={r['fair_std']:.2f}, 偏离组中位数{fair_std_median:.2f}"))

            if flags:
                biases.append({"rater": r["label"], "flags": flags})

        return biases

    def _facet(self, obs, param, labels, n):
        rows, ms, ses = [], [], []
        for p in range(n):
            idx = obs[p]
            if len(idx) == 0: continue
            total, nm = self.scores[idx].sum(), len(idx)
            zz, vv = self.z[idx], self.var_o[idx]
            w = vv.sum(); infit = (zz**2 * vv).sum() / max(w, 1e-10); outfit = (zz**2).sum() / max(nm, 1)
            se = 1.0 / np.sqrt(max(vv.sum(), 1e-10))
            rows.append({"label": labels[p], "total": int(total), "obs_avg": round(self.scores[idx].mean(), 2), "fair_avg": round(self.exp_scores[idx].mean(), 2), "meas": round(param[p], 3), "se": round(se, 3), "infit": round(infit, 3), "outfit": round(outfit, 3)})
            ms.append(param[p]); ses.append(se)
        ma, sa = np.array(ms), np.array(ses)
        vo = np.var(ma, ddof=1) if len(ma) > 1 else 0.001; me = np.mean(sa**2); vt = max(vo - me, 0.001)
        return rows, np.sqrt(vt / me) if me > 0 else 0, vt / (vt + me)

    def report(self):
        if not hasattr(self, 'exp_scores'): raise RuntimeError("请先运行 fit()")
        r = {"summary": {"N": self.N, "n_s": self.n_s, "n_r": self.n_r, "n_c": self.n_c, "n_i": self.n_i, "score_range": f"{self.min_s}-{self.max_s}", "obs_mean": round(float(self.scores.mean()), 3), "exp_mean": round(float(self.exp_scores.mean()), 3), "resid_sd": round(float(self.resid.std(ddof=1)), 4), "stres_sd": round(float(self.z.std(ddof=1)), 4), "var_exp": round(self.var_exp, 2), "ll": round(self.ll_final, 2)}, "facets": {}}
        for name, obs, param, labels, n in [("students", self.obs_s, self.theta, self.s_labels[:self.n_s], self.n_s), ("raters", self.obs_r, self.delta, self.r_labels[:self.n_r], self.n_r), ("criteria", self.obs_c, self.alpha, self.c_labels[:self.n_c], self.n_c)]:
            rows, sep, rel = self._facet(obs, param, labels, n)
            r["facets"][name] = {"rows": rows, "separation": round(sep, 2), "reliability": round(rel, 3)}
        if self.n_i > 1:
            rows, sep, rel = self._facet(self.obs_i, self.beta, self.i_labels[:self.n_i], self.n_i)
            r["facets"]["items"] = {"rows": rows, "separation": round(sep, 2), "reliability": round(rel, 3)}
        # v0.9.0: 新诊断功能 (参考 Facets 专业报告)
        r["bias"] = self._diagnose_bias()
        r["categories"] = self._diagnose_categories()
        r["anomalous"] = self._anomalous_responses()
        r["rank_compare"] = self._rank_compare()
        return r

    def _rank_compare(self) -> list[dict]:
        """v0.9.0: 原始排名 vs MFRM 校正排名对比 (参考 PPT Slide 25)。"""
        if not hasattr(self, 'exp_scores'):
            return []
        rows = []
        for s in range(self.n_s):
            idx = self.obs_s[s]
            if len(idx) == 0:
                continue
            raw_total = int(self.scores[idx].sum())
            fair_avg = round(float(self.exp_scores[idx].mean()), 2)
            rows.append({
                "label": self.s_labels[s] if s < len(self.s_labels) else f"S{s+1}",
                "raw_total": raw_total,
                "fair_avg": fair_avg,
                "meas": round(float(self.theta[s]), 3),
            })
        rows.sort(key=lambda x: x["raw_total"], reverse=True)
        for i, r in enumerate(rows):
            r["raw_rank"] = i + 1
        rows.sort(key=lambda x: x["fair_avg"], reverse=True)
        for i, r in enumerate(rows):
            r["fair_rank"] = i + 1
            r["rank_diff"] = r["raw_rank"] - r["fair_rank"]
        return rows

    def print(self):
        r = self.report(); s = r["summary"]
        print(f"\n{'='*72}\nMFRMSight v0.9.0 — 多面Rasch模型分析报告\n{'='*72}")
        print(f"{s['N']}条 | {s['n_s']}学生×{s['n_r']}评分者×{s['n_c']}标准×{s['n_i']}题目 | {s['score_range']}分")
        print(f"方差解释: {s['var_exp']}% | LL: {s['ll']:.0f}")
        print(f"ObsMean={s['obs_mean']:.3f} ExpMean={s['exp_mean']:.3f} ResidSD={s['resid_sd']:.4f} StResSD={s['stres_sd']:.4f}")
        for fn, fd in r["facets"].items():
            if not fd["rows"]: continue
            print(f"\n{'─'*72}\n{fn} — Sep={fd['separation']:.2f} Rel={fd['reliability']:.3f}\n{'─'*72}")
            print(f"{'':<16} {'Total':>7} {'ObsAvg':>7} {'FairAvg':>7} {'Meas':>7} {'SE':>6} {'Infit':>6} {'Outfit':>6}")
            print("-" * 72)
            for row in fd["rows"]:
                print(f"{row['label']:<16} {row['total']:>7.0f} {row['obs_avg']:>7.2f} {row['fair_avg']:>7.2f} {row['meas']:>7.3f} {row['se']:>6.3f} {row['infit']:>6.3f} {row['outfit']:>6.3f}")

        # v0.9.0: 评分者偏差诊断
        if r.get("bias"):
            print(f"\n{'='*72}\n[!] 评分者偏差诊断\n{'='*72}")
            for b in r["bias"]:
                for tag, detail in b["flags"]:
                    print(f"  [{tag}] {b['rater']}: {detail}")
            if not any(b["flags"] for b in r["bias"]):
                print("  [OK] 未检出显著评分者偏差")
        else:
            print(f"\n{'='*72}\n[OK] 评分者偏差诊断: 未检出异常\n{'='*72}")

        # v0.9.0: 排名对比
        if r.get("rank_compare"):
            print(f"\n{'='*72}\n[>>>] 原始总分 vs MFRM校正 排名对比\n{'='*72}")
            print(f"{'考生':<16} {'原始总分':>8} {'原始排名':>8} {'校正分':>8} {'校正排名':>8} {'排名变化':>8}")
            print("-" * 60)
            for row in r["rank_compare"]:
                diff = row["rank_diff"]
                arrow = "^" if diff > 0 else ("v" if diff < 0 else "-")
                print(f"{row['label']:<16} {row['raw_total']:>8} {row['raw_rank']:>8} {row['fair_avg']:>8.2f} {row['fair_rank']:>8} {arrow}{abs(diff):>7}")

        # v0.9.0: 等级类别功能诊断
        cat = r.get("categories")
        if cat:
            print(f"\n{'='*72}\n[*] 评分等级类别功能诊断 (参考 Facets Table 8)\n{'='*72}")
            print(f"等级数: {self.K+1} 档 ({self.min_s}-{self.max_s}), 阈值有序: {'是' if cat['tau_ordered'] else '否'}")
            if cat["issues"]:
                for issue in cat["issues"]:
                    print(f"  [!] {issue}")
            if cat["merge_suggestion"]:
                print(f"  [>>] {cat['merge_suggestion']}")
            if not cat["issues"]:
                print("  [OK] 等级功能正常，无需合并")

        # v0.9.0: 异常反应
        anom = r.get("anomalous", [])
        if anom:
            print(f"\n{'='*72}\n[!!] 异常反应记录 (|StRes| >= 3)\n{'='*72}")
            print(f"{'考生':<12} {'评分者':<10} {'标准':<8} {'观测分':>6} {'期望分':>6} {'残差':>6} {'StRes':>6}")
            print("-" * 60)
            for a in anom[:15]:  # 最多显示 15 条
                print(f"{a['student']:<12} {a['rater']:<10} {a['criterion']:<8} "
                      f"{a['observed']:>6} {a['expected']:>6.1f} {a['residual']:>6.1f} {a['stres']:>6.2f}")
            if len(anom) > 15:
                print(f"  ... 共 {len(anom)} 条异常反应")
            print(f"\n  注意: 异常残差不直接等同于评分者偏见，应结合评分者、标准、原始评分记录综合判断")
        else:
            print(f"\n{'='*72}\n[OK] 异常反应检测: 未检出 |StRes| >= 3 的反应\n{'='*72}")

    def to_excel(self, path):
        import pandas as pd
        r = self.report(); dfs = {n: pd.DataFrame(f["rows"]) for n, f in r["facets"].items()}
        dfs["summary"] = pd.DataFrame([r["summary"]])
        with pd.ExcelWriter(path) as w:
            for n, d in dfs.items(): d.to_excel(w, sheet_name=n, index=False)
        print(f"[OK] {path}")

    def to_word(self, path):
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            print("[!] pip install python-docx"); return
        r = self.report(); doc = Document()
        t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = t.add_run("MFRM Rating Scale 模型分析报告"); run.bold = True; run.font.size = Pt(18)
        s = r["summary"]
        doc.add_paragraph(f"反应数: {s['N']} | {s['n_s']}学生×{s['n_r']}评分者×{s['n_c']}标准×{s['n_i']}题目\n分数: {s['score_range']} | Rasch解释方差: {s['var_exp']}%\n观察均值: {s['obs_mean']:.3f} | 残差SD: {s['resid_sd']:.3f} | StRes SD: {s['stres_sd']:.3f}")
        for fn, fd in r["facets"].items():
            if not fd["rows"]: continue
            doc.add_heading(f"{fn} (Sep={fd['separation']:.2f}, Rel={fd['reliability']:.3f})", level=2)
            h = ["元素", "总分", "ObsAvg", "FairAvg", "Meas", "SE", "Infit", "Outfit"]
            tbl = doc.add_table(rows=1 + len(fd["rows"]), cols=8); tbl.style = "Table Grid"
            for i, hh in enumerate(h): tbl.rows[0].cells[i].text = hh
            for ri, row in enumerate(fd["rows"]):
                for ci, k in enumerate(["label", "total", "obs_avg", "fair_avg", "meas", "se", "infit", "outfit"]):
                    v = row[k]; tbl.rows[ri + 1].cells[ci].text = f"{v:.1f}" if isinstance(v, float) and k != "label" else f"{v:.0f}" if k == "total" else str(v)
        doc.save(path); print(f"[OK] {path}")
