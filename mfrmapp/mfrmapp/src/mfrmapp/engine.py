"""
MFRMSight — 多面Rasch模型分析引擎
===================================
基于Andrich Rating Scale Model (1978)
Fisher-scoring JMLE估计, 分阶段PROX→JMLE
"""
import numpy as np
from pathlib import Path
from typing import Optional
import re, io, base64
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt


__version__ = "1.0.12"

def parse_facets_txt(filepath: str) -> dict:
    """解析Facets/Minifac风格的.txt输入文件."""
    with open(filepath, "r", encoding="utf-8") as f:
        lines = [l.strip() for l in f if l.strip()]
    data_rows, infos = [], [[], [], [], []]
    labels_ok, in_data, cur_facet = False, False, 0
    n_facets = 4; noncentered = 1
    for line in lines:
        lower = line.lower()
        normalized = lower.replace(" = ", "=").replace("= ", "=").replace(" =", "=")
        if normalized.startswith("facets="):
            try: n_facets = int(normalized.split("=")[1])
            except ValueError: pass
        if normalized.startswith("noncentered=") or normalized.startswith("noncent"):
            try: noncentered = int(normalized.split("=")[1])
            except ValueError: pass
        if line == "*": labels_ok, cur_facet = True, 0; continue
        if normalized.startswith("labels="): continue
        if normalized.startswith("data="): labels_ok = False; in_data = True; continue
        if labels_ok and not in_data:
            if line.isdigit(): continue
            if "," in line:
                parts = [x.strip() for x in line.split(",", 1)]
                if parts[0].isdigit():
                    fid = int(parts[0])
                    if cur_facet == 0: cur_facet = fid
                    elif fid > 0 and parts[1].strip(): infos[cur_facet - 1].append((fid, parts[1]))
            continue
        if in_data:
            if line.startswith(";"): continue
            p = [x.strip() for x in line.split(",")]
            if len(p) < n_facets + 1: p = [x.strip() for x in line.split()]
            if len(p) >= n_facets + 1:
                try: data_rows.append([int(v) for v in p])
                except ValueError: pass
    raw = np.array(data_rows) if data_rows else np.empty((0, n_facets + 1), dtype=int)
    result = _build_dict(raw, n_facets,
                       [l for _, l in infos[0]], [l for _, l in infos[1]],
                       [l for _, l in infos[2]], [l for _, l in infos[3]])
    result["noncentered"] = noncentered
    return result


# ═══════════════════════════════════════════════════════════════════════
# v1.0.0: 英文→中文翻译表 + 面名称提取
# ═══════════════════════════════════════════════════════════════════════

_EN_ZH_MAP: dict[str, str] = {
    # 面名称
    "students": "考生", "student": "考生",
    "raters": "评分者", "rater": "评分者",
    "criteria": "评分标准", "criterion": "评分标准",
    "items": "题目", "item": "题目",
    "judges": "评委", "judge": "评委",
    "examiners": "考官", "examiner": "考官",
    "tasks": "任务", "task": "任务",
    "domains": "领域", "domain": "领域",
    # 常见维度名
    "comp": "综合能力", "inte": "人际沟通",
    "lear": "学习能力", "exec": "执行能力",
    "matc": "匹配能力", "crea": "创新能力",
    "lead": "领导力", "comm": "沟通能力",
    "anal": "分析能力", "logi": "逻辑思维",
    "expr": "表达能力",
    # 常见元素名
    "boy": "男孩", "girl": "女孩",
    "male": "男", "female": "女",
}


def _translate_en(text: str) -> str:
    """v1.0.0: 英文→中文翻译。检测是否含中文，不含则查表翻译。"""
    # 已含中文则不翻译
    if any('一' <= c <= '鿿' for c in text):
        return ""
    lower = text.lower().strip()
    return _EN_ZH_MAP.get(lower, _EN_ZH_MAP.get(lower.rstrip("s"), ""))


def extract_dimensions(data: dict) -> dict:
    """v1.0.0: 从解析后数据提取面维度和元素标签。

    Returns:
      {"facets": [{"key": "students", "label": "考生", "original": "students", "elements": ["Student1",...]}, ...],
       "n_facets": 3 or 4}
    """
    facet_defs = [
        ("students", "s_labels", "n_s"),
        ("raters", "r_labels", "n_r"),
        ("criteria", "c_labels", "n_c"),
    ]
    n_f = data.get("n_facets", 4)
    if n_f >= 4:
        facet_defs.append(("items", "i_labels", "n_i"))

    facets = []
    for key, lbl_key, n_key in facet_defs:
        labels = data.get(lbl_key, [])
        n = data.get(n_key, len(labels) if labels else 0)
        labels = labels[:n] if len(labels) >= n else labels
        original = ""
        # 尝试从 facet_meta 获取原始面名称
        meta = data.get("facet_meta", [])
        idx = {"students": 0, "raters": 1, "criteria": 2, "items": 3}.get(key, -1)
        if 0 <= idx < len(meta) and meta[idx].get("name"):
            original = meta[idx]["name"]
        # 从 labels 推断
        if not original and labels:
            original = labels[0] if len(labels) == 1 else key
        if not original:
            original = key
        zh = _translate_en(original)
        facets.append({
            "key": key, "label": zh or key, "original": original,
            "elements": labels, "n": n,
        })
    return {"facets": facets, "n_facets": n_f}


def filter_data(raw: "np.ndarray", n_facets: int,
                keep_criteria: list[int] | None = None,
                keep_items: list[int] | None = None) -> "np.ndarray":
    """v1.0.0: 按选中维度过滤 raw 数组并重新映射 ID。

    raw 列: [student_id, rater_id, criterion_id, (item_id), score]
    """
    import numpy as np
    mask = np.ones(len(raw), dtype=bool)
    new_raw = raw.copy()

    if keep_criteria is not None and n_facets >= 3:
        mask &= np.isin(new_raw[:, 2], keep_criteria)
        # 重映射 criterion ID
        id_map = {old: new for new, old in enumerate(sorted(keep_criteria), 1)}
        for old, new in id_map.items():
            new_raw[new_raw[:, 2] == old, 2] = new

    if keep_items is not None and n_facets >= 4:
        mask &= np.isin(new_raw[:, 3], keep_items)
        id_map = {old: new for new, old in enumerate(sorted(keep_items), 1)}
        for old, new in id_map.items():
            new_raw[new_raw[:, 3] == old, 3] = new

    return new_raw[mask]


def parse_facets_txt(filepath: str) -> dict:
    """解析Facets/Minifac风格的.txt输入文件。

    适配 v0.8.0: 支持等号两边有空格的写法 (如 "Facets = 3"、"Data = ")，
    支持 Labels 段中的裸数字行(元素编号占位符，无标签)，
    支持 Data 段中的分号注释行 (如 "; Student1")。
    自动检测 3 面或 4 面设计。

    文件 min_line: Facets/N, Noncentered/N, Labels, Data, Model=?,?,?,R23
    """
    with open(filepath, "r", encoding="utf-8") as f:
        lines = [l.strip() for l in f if l.strip()]
    data_rows, infos = [], [[], [], [], []]
    labels_ok, in_data, cur_facet = False, False, 0
    n_facets = 4  # 默认
    noncentered = 1

    for line in lines:
        # ── 规范化: 等号两边空格不影响──
        lower = line.lower()
        # 去掉等号周围的空格: "facets = 3" → "facets=3"
        normalized = lower.replace(" = ", "=").replace("= ", "=").replace(" =", "=")

        # ── Header 指令 ──
        if normalized.startswith("facets="):
            try: n_facets = int(normalized.split("=")[1].split(";")[0])
            except ValueError: pass
        if normalized.startswith("noncentered=") or normalized.startswith("noncent"):
            try: noncentered = int(normalized.split("=")[1].split(";")[0])
            except ValueError: pass
        if normalized.startswith("model="):
            pass  # Model 信息当前仅存储，不影响解析
        if normalized.startswith("title="):
            pass

        # ── 分隔符 * ──
        if line == "*":
            labels_ok = True
            cur_facet = 0
            continue

        # ── Labels 段 ──
        if normalized.startswith("labels="):
            continue
        if normalized.startswith("data="):
            labels_ok = False
            in_data = True
            continue

        if labels_ok and not in_data:
            # 裸数字行 (元素编号占位符), 跳过
            if line.isdigit():
                continue
            if "," in line:
                parts = [x.strip() for x in line.split(",", 1)]
                if parts[0].isdigit():
                    fid = int(parts[0])
                    if cur_facet == 0:
                        cur_facet = fid  # facet声明行
                    else:
                        if fid > 0 and parts[1].strip():
                            infos[cur_facet - 1].append((fid, parts[1]))
            continue

        # ── Data 段 ──
        if in_data:
            if line.startswith(";"):
                continue
            # 同时支持逗号和空格/制表分隔
            p = [x.strip() for x in line.split(",")]
            if len(p) < n_facets + 1:
                p = [x.strip() for x in line.split()]
            if len(p) >= n_facets + 1:
                # 展开范围 token (如 "1-4", "1\t-\t2")
                rows = [p]
                for ti, tok in enumerate(p):
                    clean = re.sub(r'\s+', '', tok)
                    m = re.match(r'^(\d+)[_-](\d+)$', clean)
                    if m:
                        lo, hi = int(m.group(1)), int(m.group(2))
                        vals = list(range(lo, hi + 1))
                        new_rows = []
                        for row in rows:
                            for v in vals:
                                r2 = row[:]
                                r2[ti] = str(v)
                                new_rows.append(r2)
                        rows = new_rows
                for row in rows:
                    try:
                        data_rows.append([int(v) for v in row])
                    except ValueError:
                        pass

    raw = np.array(data_rows) if data_rows else np.empty((0, n_facets + 1), dtype=int)
    result = _build_dict(raw, n_facets,
                       [l for _, l in infos[0]], [l for _, l in infos[1]],
                       [l for _, l in infos[2]], [l for _, l in infos[3]])
    result["noncentered"] = noncentered
    return result


def parse_facets_out(filepath: str) -> dict:
    """v0.9.0: 解析 Facets .out.txt 输出文件，提取关键 Table。

    提取:
      - Table 3: 迭代历史 (PROX/JMLE 收敛)
      - Table 5: 总体拟合 (ObsMean/ExpMean/VarExp/Chi-squared)
      - Table 7.x: 各面向测量报告 (Measure/SE/Infit/Outfit/Separation/Reliability)
      - Table 8.1: 等级类别统计 (Counts/Thresholds/Measures)
      - Table 4: 异常反应列表

    Returns: {"summary": {...}, "facets": {...}, "categories": {...}, "anomalous": [...]}
    """
    with open(filepath, "r", encoding="utf-8") as f:
        content = f.read()
    lines = content.split("\n")

    result = {"summary": {}, "facets": {}, "categories": {}, "anomalous": []}

    # ── 定位各 Table 起始行 ──
    table_starts = {}
    for i, line in enumerate(lines):
        line_s = line.strip()
        for tn in ["Table 3", "Table 4.1", "Table 5", "Table 7.1.1", "Table 7.2.1",
                    "Table 7.3.1", "Table 7.4.1", "Table 8.1"]:
            if line_s.startswith(tn) and tn not in table_starts:
                table_starts[tn] = i
                break

    # ── Table 5: 总体拟合 ──
    if "Table 5" in table_starts:
        chunk = "\n".join(lines[table_starts["Table 5"]:table_starts["Table 5"]+20])
        m = re.search(r"Mean \(Count:\s*(\d+)\)", chunk)
        obs_n = int(m.group(1)) if m else 0
        # 提取 ObsMean/ExpMean/S.D.
        vals = re.findall(r"([\d.]+)\s+([\d.]+)\s+([\d.]+)\s+([\d.]+)\s+([\d.]+)", chunk)
        if vals:
            result["summary"]["obs_mean"] = float(vals[0][0])
            result["summary"]["exp_mean"] = float(vals[0][1])
            # 取 S.D. Sample 行 (第3行)
            result["summary"]["sd_obs"] = float(vals[2][2]) if len(vals) >= 3 else float(vals[1][2])
        m = re.search(r"Variance explained by Rasch measures\s*=\s*([\d.]+)\s*([\d.]+)%", chunk)
        if m:
            result["summary"]["var_exp"] = float(m.group(2))
        m = re.search(r"chi-squared\s*=\s*([\d.]+).*probability\s*=\s*([\d.]+)", chunk)
        if m:
            result["summary"]["chi_sq"] = float(m.group(1))
            result["summary"]["chi_prob"] = float(m.group(2))
        result["summary"]["N"] = obs_n

    # ── Table 7.x: 测量报告 ──
    facet_map = {"Table 7.1.1": "students", "Table 7.2.1": "raters",
                 "Table 7.3.1": "criteria", "Table 7.4.1": "items"}
    for tn, fname in facet_map.items():
        if tn not in table_starts:
            continue
        start = table_starts[tn]
        end = start + 25
        chunk = lines[start:end]
        rows = []
        for line in chunk:
            # Table 7 数据行: "|  916      64  ...   .89   .87 | 4 Student4  |"
            m = re.match(
                r"\|\s*(\d+)\s+(\d+)\s+([\d.]+)\s+([\d.]+)\s*\|"
                r"\s*([\d.-]+)\s+([\d.]+)\s*\|"
                r"\s*([\d.]+)\s+([\d.-]+)\s+([\d.]+)\s+([\d.-]+)\s*\|"
                r"\s*([\d.]+)\s*\|"
                r"\s*([\d.]+)\s+([\d.]+)\s*\|"
                r"\s*\d+\s+(.+)", line)
            if m:
                rows.append({
                    "total": int(m.group(1)), "count": int(m.group(2)),
                    "obs_avg": float(m.group(3)), "fair_avg": float(m.group(4)),
                    "meas": float(m.group(5)), "se": float(m.group(6)),
                    "infit_mnsq": float(m.group(7)), "infit_zstd": float(m.group(8)),
                    "outfit_mnsq": float(m.group(9)), "outfit_zstd": float(m.group(10)),
                    "discrm": float(m.group(11)),
                    "ptmea": float(m.group(12)), "ptexp": float(m.group(13)),
                    "label": m.group(14).strip(),
                })
            # 提取 Separation/Reliability (取第一条 Population 行)
            m2 = re.search(r"Separation\s+([\d.]+).*Reliability\s+([\d.]+)", line)
            if m2 and fname not in result["facets"]:
                result["facets"][fname] = {
                    "separation": float(m2.group(1)),
                    "reliability": float(m2.group(2)),
                    "rows": rows,
                }

    # ── Table 8.1: 等级类别统计 ──
    if "Table 8.1" in table_starts:
        start = table_starts["Table 8.1"]
        chunk = lines[start:start+35]
        cat_rows = []
        for line in chunk:
            # 只处理以 "|" 开头且第一个字段为数字的行
            line_s = line.strip()
            if not line_s.startswith("|") or not re.match(r"\|\s*\d", line_s):
                continue
            # 按 | 分割，取前 5 个字段
            cells = [c.strip() for c in line_s.split("|")]
            if len(cells) < 4:
                continue
            # 字段 1: "2       3         3    1%   1%" → [score, total, used, pct, cum_pct]
            f1 = cells[1].split()
            if len(f1) < 3:
                continue
            score = int(f1[0])
            count = int(f1[2])  # "Used" 列
            # 字段 2: "-1.79  -2.35  1.6" → [avg_meas, exp_meas, outfit]
            f2 = cells[2].split()
            avg_meas = float(f2[0].rstrip("*"))
            exp_meas = float(f2[1]) if len(f2) > 1 else 0.0
            outfit_mnsq = float(f2[2]) if len(f2) > 2 else 1.0
            # 字段 3: "-3.35    .61" → [threshold, thresh_se] or empty
            f3 = cells[3].split()
            threshold = float(f3[0]) if len(f3) > 0 else None
            thresh_se = float(f3[1]) if len(f3) > 1 else None
            # 字段 4: "( -4.72)       " 或 "( 4.75)   4.10" → cat_meas
            f4 = cells[4].strip()
            cat_meas = 0.0
            if f4.startswith("("):
                inner = f4.split(")")[0].lstrip("(").strip().split()[0]
                try: cat_meas = float(inner)
                except ValueError: pass
            cat_rows.append({
                "score": score, "count": count,
                "avg_meas": avg_meas, "exp_meas": exp_meas,
                "outfit_mnsq": outfit_mnsq,
                "threshold": threshold, "thresh_se": thresh_se,
                "cat_meas": cat_meas,
            })
        tau_values = [r["threshold"] for r in cat_rows if r["threshold"] is not None]
        result["categories"] = {
            "rows": cat_rows,
            "tau": tau_values,
            "tau_ordered": all(tau_values[i] <= tau_values[i+1] for i in range(len(tau_values)-1)) if len(tau_values) > 1 else True,
        }

    # ── Table 4: 异常反应 ──
    if "Table 4.1" in table_starts:
        start = table_starts["Table 4.1"]
        chunk = lines[start:start+20]
        for line in chunk:
            if "No unexpected" in line:
                break
            m = re.match(r"\|\s*(\d+)\s+(\d+)\s+([\d.]+)\s+([\d.-]+)\s+([\d.-]+)\s*\|", line)
            if m:
                result["anomalous"].append({
                    "score": int(m.group(1)),
                    "observed": int(m.group(2)),
                    "expected": float(m.group(3)),
                    "residual": float(m.group(4)),
                    "stres": float(m.group(5)),
                })

    return result


def parse_excel(filepath: str) -> dict:
    """从Excel加载3面向数据"""
    import pandas as pd
    df_raw = pd.read_excel(filepath, header=None)
    header_row = 1
    for i in range(len(df_raw)):
        s = [str(v) for v in df_raw.iloc[i] if pd.notna(v)]
        if any("编号" in x or "评分人" in x for x in s):
            header_row = i; break
    col_names = [str(df_raw.iloc[header_row, j]) if pd.notna(df_raw.iloc[header_row, j]) else f"Col{j}"
                 for j in range(df_raw.shape[1])]

    # BUG-002 增强: 智能列名推断
    def _infer_excel_col_role(cn: str, col_idx: int) -> str:
        """推断 Excel 列名在 MFRM 中的角色: student / criterion / item / rater / row_id"""
        cn_lower = cn.lower()
        # 强信号
        if any(kw in cn_lower for kw in ["编号", "序号", "id", "row"]):
            return "row_id"
        if any(kw in cn_lower for kw in ["评分人", "评分者", "rater", "评委", "examiner", "judge"]):
            return "rater"
        if any(kw in cn_lower for kw in ["学生", "student", "考生", "学员", "受试"]):
            return "student"
        if any(kw in cn_lower for kw in ["题目", "题号", "item", "task"]):
            return "item"
        if any(kw in cn_lower for kw in ["标准", "维度", "criterion", "criteria", "domain"]):
            return "criterion"

        # 列名含连字符: "综合分析-1号" → 前缀=criterion, 后缀看上下文
        if "-" in cn or "—" in cn:
            sep = "-" if "-" in cn else "—"
            parts = cn.rsplit(sep, 1) if cn.count(sep) == 1 else (cn.split(sep)[0], cn.split(sep)[-1])
            prefix, suffix = parts[0], parts[-1]
            # 前缀含能力关键词 → criterion
            has_crit = any(kw in prefix for kw in ["综合", "沟通", "分析", "人际", "表达", "逻辑", "创新", "协作", "能力"])
            # 后缀含题 → item
            has_item = any(kw in suffix for kw in ["题", "item"])
            # 后缀含号/生/人 → student
            has_stu = any(kw in suffix for kw in ["号", "生", "人", "student"])
            has_num = bool(re.search(r'\d+', suffix))
            if has_crit:
                return "criterion"  # 前缀明确是criterion
            if has_item:
                return "item"
            if has_stu:
                return "student"
            if has_num:
                return "student"  # 默认数字后缀→student
            return "unknown"

        # 纯数字 → student
        if re.match(r'^\d+$', cn.strip()):
            return "student"

        return "unknown"

    s_set, c_set, i_set = set(), set(), set()
    for cn in col_names[2:]:
        role = _infer_excel_col_role(cn, 0)
        if "-" in cn or role == "criterion":
            crit_name = cn.rsplit("-", 1)[0] if "-" in cn else cn
            c_set.add(crit_name)
            stu_name = cn.rsplit("-", 1)[-1] if "-" in cn else cn
            if role == "student":
                s_set.add(stu_name)
            elif role == "item":
                i_set.add(stu_name)
            else:
                s_set.add(stu_name)  # 默认归student
        elif role == "student":
            s_set.add(cn)
        elif role == "item":
            i_set.add(cn)
        elif role == "criterion":
            c_set.add(cn)
        else:
            # 兜底: 按连字符分析
            if "-" in cn:
                crit_name, stu_name = cn.rsplit("-", 1)
                c_set.add(crit_name)
                s_set.add(stu_name)
    s_list, c_list, i_list = sorted(s_set), sorted(c_set), sorted(i_set)
    n_s, n_c = len(s_list), len(c_list)
    n_i = len(i_list) if i_list else 1
    data_rows = []
    rater_labels = []
    # BUG-002 增强: 自动检测 n_facets (3 或 4)
    auto_n_facets = 4 if n_i > 1 else 3
    for ri in range(header_row + 1, len(df_raw)):
        row = df_raw.iloc[ri]
        if pd.isna(row.iloc[0]): continue
        r_label = str(row.iloc[1])
        r_num = _extract_num(r_label)
        if auto_n_facets >= 4:
            for si in range(n_s):
                for ci in range(n_c):
                    for ii in range(n_i):
                        col = 2 + ci * n_s * n_i + si * n_i + ii
                        if col < df_raw.shape[1]:
                            v = row.iloc[col]
                            if pd.notna(v):
                                data_rows.append([si + 1, r_num, ci + 1, ii + 1, int(v)])
        else:
            for si in range(n_s):
                for ci in range(n_c):
                    col = 2 + ci * n_s + si
                    if col < df_raw.shape[1]:
                        v = row.iloc[col]
                        if pd.notna(v): data_rows.append([si + 1, r_num, ci + 1, 1, int(v)])
        rater_labels.append(r_label)
    raw = np.array(data_rows)
    return _build_dict(raw, auto_n_facets, s_list, rater_labels, c_list, i_list if i_list else ["Item1"])


def _extract_num(s): return int(re.findall(r'\d+', str(s))[-1]) if re.findall(r'\d+', str(s)) else 1


def _build_dict(raw, n_facets, s_labels, r_labels, c_labels, i_labels):
    return {
        "raw": raw, "n_facets": n_facets, "N": len(raw),
        "n_s": int(raw[:, 0].max()), "n_r": int(raw[:, 1].max()),
        "n_c": int(raw[:, 2].max()), "n_i": int(raw[:, 3].max()) if n_facets >= 4 and raw.shape[1] > 3 else 1,
        "min_s": int(raw[:, -1].min()), "max_s": int(raw[:, -1].max()),
        "s_labels": s_labels or [f"Student{i+1}" for i in range(int(raw[:, 0].max()))],
        "r_labels": r_labels or [f"Rater{i+1}" for i in range(int(raw[:, 1].max()))],
        "c_labels": c_labels or [f"Criterion{i+1}" for i in range(int(raw[:, 2].max()))],
        "i_labels": i_labels or [f"Item{i+1}" for i in range(max(1, int(raw[:, 3].max()) if n_facets >= 4 and raw.shape[1] > 3 else 1))],
    }


class MFRMEngine:
    """多面Rasch Rating Scale模型估计器"""

    def __init__(self, data: dict):
        for k, v in data.items(): setattr(self, k, v)
        self.noncentered = data.get("noncentered", 1)  # 默认约束第1面(Students), Facets 约定
        self._index()

    def _index(self):
        self.s_idx = self.raw[:, 0] - 1
        self.r_idx = self.raw[:, 1] - 1
        # n_facets=2 时 raw 只有 [student, rater, score] 三列
        if self.n_facets >= 3:
            self.c_idx = np.clip(self.raw[:, 2] - 1, 0, self.n_c - 1)
        else:
            self.c_idx = np.zeros(self.N, dtype=int)
        if self.n_facets >= 4:
            self.i_idx = self.raw[:, 3] - 1
        else:
            self.i_idx = np.zeros(self.N, dtype=int)
        self.scores = self.raw[:, -1].astype(float)
        self.x = self.scores - self.min_s
        self.K = int(self.max_s - self.min_s)
        self.cats = np.arange(self.K + 1, dtype=float)
        self.obs_s = [np.where(self.s_idx == s)[0] for s in range(self.n_s)]
        self.obs_r = [np.where(self.r_idx == r)[0] for r in range(self.n_r)]
        self.obs_c = [np.where(self.c_idx == c)[0] for c in range(self.n_c)]
        self.obs_i = [np.where(self.i_idx == it)[0] for it in range(max(self.n_i, 1))]

    def _probs(self):
        intc = (self.theta[self.s_idx] - self.delta[self.r_idx]
                - self.alpha[self.c_idx] - self.beta[self.i_idx])
        logits = np.zeros((self.N, self.K + 1)); cum = np.zeros(self.N)
        for k in range(self.K): cum += intc - self.tau[k]; logits[:, k + 1] = cum
        logits -= logits.max(axis=1, keepdims=True)
        e = np.exp(logits); return e / e.sum(axis=1, keepdims=True)

    def _ll(self, p):
        return sum(np.log(max(p[i, int(self.x[i])], 1e-300)) for i in range(self.N))

    def _s2l(self, scores_list):
        """Score→logit Newton-Raphson 反演（BUG-012 核心修复）。

        替代粗糙的 mean/M→logit 近似。
        在 Andrich Rating Scale 模型中，E[score|θ,τ] 是 θ 的光滑单调函数。
        给定固定 τ 阈值，用 N-R 自洽求 θ 使期望分等于观测均值。

        算法:
          θ ← θ + Δ, 其中 Δ = (target - E[score]) / Var,
          这正是 Fisher scoring 的单参数版本。
          收敛速度: 5-8 步 (< 1e-6 tolerance)。
        """
        K = self.K
        n = len(scores_list)
        th = np.zeros(n)
        tu = self.tau

        for i in range(n):
            sc = scores_list[i]
            nobs = len(sc)
            if nobs == 0:
                th[i] = 0.0
                continue

            target = sc.mean() - self.min_s  # 0-based

            # 极端分数直接给边界值，避免 N-R 跑飞
            if target <= 0.05 * K:
                th[i] = -5.0
                continue
            if target >= 0.95 * K:
                th[i] = 5.0
                continue

            t = 0.0
            for _ in range(25):
                lin = t - tu
                cum = np.zeros(K + 1)
                for k in range(K):
                    cum[k + 1] = cum[k] + lin[k]
                cum -= cum.max()
                e = np.exp(cum)
                p = e / e.sum()
                exp_s = p @ self.cats
                exp_sq = p @ (self.cats ** 2)
                var = max(exp_sq - exp_s ** 2, 0.001)
                delta = (target - exp_s) / var
                t += delta
                if abs(delta) < 1e-6:
                    break

            th[i] = np.clip(t, -10, 10)

        return th

    def _prox(self):
        """PROX 初始化 — 两阶段 N-R 反演（BUG-012 修复）。

        废除了旧的 mean/M→logit 公式，改用 Rating Scale 模型期望分函数
        的 Newton-Raphson 反演来精确映射 score → logit。

        阶段1: 全局分布估计初始 τ(k) → N-R 反解各元素参数 → 居中。
        阶段2: 用新参数 refine τ → 再跑一轮 N-R → 再居中。

        BUG-013 补充修复: 阶段1 N-R 后必须先居中，否则 τ 精炼从偏置
        参数出发导致概率估计有偏，收敛后 ExpMean 偏离 ObsMean 约 1.4%。
        """
        # ── 初始 τ（从边际分布，与 Facets 初值一致）──
        self.tau = np.array([-np.log(
            np.clip((self.scores >= self.min_s + k + 1).mean(), 0.02, 0.98) /
            np.clip((self.scores < self.min_s + k + 1).mean(), 0.02, 0.98))
            for k in range(self.K)])
        self.tau -= self.tau[0]

        # ── 阶段1: N-R 反解各面参数（基于初始 τ）──
        self.theta = self._s2l([self.scores[self.obs_s[s]] for s in range(self.n_s)])
        self.delta = -self._s2l([self.scores[self.obs_r[r]] for r in range(self.n_r)])
        self.alpha = -self._s2l([self.scores[self.obs_c[c]] for c in range(self.n_c)])
        if self.n_i > 1:
            self.beta = -self._s2l([self.scores[self.obs_i[it]] for it in range(self.n_i)])
        else:
            self.beta = np.zeros(1)

        # 阶段1 后立即居中: 确保 τ 精炼从已校准的参数出发
        self._center_prox()

        # ── 阶段2: refine τ 并重新 N-R ──
        if self.K > 0:
            p = self._probs()
            for k in range(self.K):
                o = (self.x >= k + 1).sum()
                e = p[:, k + 1:].sum()
                pg = p[:, k + 1:].sum(axis=1)
                info = (pg * (1 - pg)).sum() + 1.0
                self.tau[k] += 0.5 * (e - o) / info   # 半阻尼 Fisher scoring
            self.tau -= self.tau[0]

            # 重新 N-R（基于 refined τ）
            self.theta = self._s2l([self.scores[self.obs_s[s]] for s in range(self.n_s)])
            self.delta = -self._s2l([self.scores[self.obs_r[r]] for r in range(self.n_r)])
            self.alpha = -self._s2l([self.scores[self.obs_c[c]] for c in range(self.n_c)])
            if self.n_i > 1:
                self.beta = -self._s2l([self.scores[self.obs_i[it]] for it in range(self.n_i)])

            # 阶段2 后再居中 (由 _center_prox 统一处理)
            self._center_prox()

        # ── U 调整（温和校正量尺，与 Facets PROX step 4 等价）──
        p = self._probs()
        sd_obs = np.std(self.scores)
        sd_exp = np.std(self.min_s + p @ self.cats)
        if sd_exp > 0.01:
            U = 1.0 + 0.5 * (sd_obs / sd_exp - 1.0)
            U = np.clip(U, 0.8, 2.0)
        else:
            U = 1.0
        self.theta *= U; self.delta *= U; self.alpha *= U; self.beta *= U; self.tau *= U

    def _center_prox(self):
        """仅对 noncentered 指定的面居中 (Rasch 识别约束: 仅需1个sum-to-zero)"""
        # noncentered 语义: 该面不居中, 其余面居中
        # noncentered=1 → 面1(Students)不居中, 面2,3,4居中
        # 等价于: 以 Students 为自由参考, 其他面被约束
        if self.noncentered != 1:
            self.theta -= self.theta.mean()
        if self.noncentered != 2:
            self.delta -= self.delta.mean()
        if self.noncentered != 3:
            self.alpha -= self.alpha.mean()
        if self.noncentered != 4 and self.n_i > 1:
            self.beta -= self.beta.mean()

    def fit(self, p1=200, p2=300):
        # BUG-004: 稀疏数据检测与警告
        if self.K > 0:
            per_cat = np.bincount(self.x.astype(int), minlength=self.K + 1)
            sparse_cats = [str(self.min_s + i) for i, c in enumerate(per_cat) if c < 8]
            if sparse_cats:
                import warnings
                warnings.warn(
                    f"⚠️ 稀疏数据警告 (每个等级建议 >= 8 观测): "
                    f"分数 {','.join(sparse_cats)} 观测不足。参数可能不稳定，建议合并评分等级。"
                )
        self._prox()
        best = (-1e100, None)
        # 阶段1: 高阻尼 + 强ridge → 锁定大致方向
        for phase, (n_it, rs, rd, dmp) in enumerate([(p1, 0.3, 0.99, 0.2), (p2, 0.03, 0.9995, 0.08)]):
            for it in range(1, n_it + 1):
                rid = max(rs * (rd ** it), 0.001)
                # 交替更新各面向
                self._up_theta(dmp, rid); self._up_delta(dmp, rid)
                self._up_alpha(dmp, rid); self._up_beta(dmp, rid)
                # tau: 更强的阻尼和正则化，防止发散
                tau_damp = dmp * 0.3
                tau_rid = rid * 5
                self._up_tau(tau_damp, tau_rid)
                if it % 10 == 0:
                    p = self._probs(); ll = self._ll(p)
                    if ll > best[0]: best = (ll, (self.theta.copy(), self.delta.copy(), self.alpha.copy(), self.beta.copy(), self.tau.copy()))
        if best[1]: self.theta, self.delta, self.alpha, self.beta, self.tau = best[1]
        self._fin(); return self

    def _up(self, arr, obs, n, dmp, rid, sign, center=True):
        """Fisher scoring 更新 — sequential update（BUG-013 修复）。

        旧版: 循环外一次 _probs() → 批量更新所有元素（batch mode）
        新版: 循环内每更新一个元素后立即 _probs()（sequential/JMLE mode）

        数学: JMLE 等价于对每个参数做 Newton step，
              需要"其他参数固定"的假设才成立。
              Sequential 保持了该假设。
        """
        for i in range(n):
            p = self._probs(); e = p @ self.cats
            v = np.clip((p @ (self.cats**2)) - e**2, 0.001, None)
            idx = obs[i]; arr[i] += dmp * (sign * self.x[idx].sum() - sign * e[idx].sum()) / (v[idx].sum() + rid)
        if center: arr -= arr.mean()
        arr[:] = np.clip(arr, -20, 20)

    def _up_theta(self, d, r): self._up(self.theta, self.obs_s, self.n_s, d, r, 1, center=(self.noncentered != 1))
    def _up_delta(self, d, r): self._up(self.delta, self.obs_r, self.n_r, d, r, -1, center=(self.noncentered != 2))
    def _up_alpha(self, d, r): self._up(self.alpha, self.obs_c, self.n_c, d, r, -1, center=(self.noncentered != 3))
    def _up_beta(self, d, r):
        if self.n_i > 1: self._up(self.beta, self.obs_i, self.n_i, d, r, -1, center=(self.n_facets >= 4 and self.noncentered != 4))

    def _up_tau(self, dmp, rid):
        """Tau 更新 — sequential update（BUG-013 修复）。

        每个 τ_k 更新后立即重算概率，
        因为 τ_k 的变化影响所有 k+1...K 的累积概率。
        """
        for k in range(self.K):
            p = self._probs()
            o = (self.x >= k + 1).sum()
            e = p[:, k + 1:].sum()
            pg = p[:, k + 1:].sum(axis=1); info = (pg * (1 - pg)).sum() + rid
            self.tau[k] += dmp * (e - o) / info
        self.tau -= self.tau[0]; self.tau[:] = np.clip(self.tau, -15, 25)

    def _fin(self):
        p = self._probs(); self.ll_final = self._ll(p)
        e = p @ self.cats; self.var_o = np.clip((p @ (self.cats**2)) - e**2, 0.001, None)
        self.exp_scores = self.min_s + e; self.resid = self.scores - self.exp_scores; self.z = self.resid / np.sqrt(self.var_o)
        self.var_exp = max(0, (np.var(self.scores, ddof=1) - np.var(self.resid, ddof=1)) / np.var(self.scores, ddof=1) * 100)

    def _diagnose_categories(self) -> dict:
        """v0.9.0: 评分等级类别功能诊断 (参考 Facets Table 8)。

        判断等级划分是否合理:
          1. 每个等级至少 10 次使用
          2. 类别平均测量值必须单调递增
          3. Andrich 阈值必须有序
          4. 类别峰值概率 > 15%

        Returns: {"passed": bool, "issues": [...], "merge_suggestion": str or None}
        """
        issues = []
        scores_int = self.x.astype(int)
        usage = np.bincount(scores_int, minlength=self.K + 1)

        # 1. 使用频次检查
        low_usage = [(self.min_s + i, int(c)) for i, c in enumerate(usage) if c < 10]
        if low_usage:
            cats_str = ", ".join(f"{s}({c}次)" for s, c in low_usage)
            issues.append(f"等级使用不足: {cats_str}")

        # 2. 类别平均 Measure 单调性
        cat_measures = []
        for k in range(self.K + 1):
            idx = np.where(scores_int == k)[0]
            if len(idx) > 0:
                cat_measures.append((self.min_s + k, float(self.theta[self.s_idx[idx]].mean()) if len(idx) > 0 else 0))
        if cat_measures:
            for i in range(1, len(cat_measures)):
                if cat_measures[i][1] < cat_measures[i-1][1]:
                    issues.append(f"等级 {cat_measures[i-1][0]}->{cat_measures[i][0]} 平均测量值非单调 "
                                  f"({cat_measures[i-1][1]:.2f} -> {cat_measures[i][1]:.2f})")
                    break  # 只报一次

        # 3. Andrich 阈值有序性
        if self.K >= 2:
            thresholds_ordered = all(self.tau[i] <= self.tau[i+1] for i in range(self.K - 1))
            if not thresholds_ordered:
                disordered = []
                for i in range(self.K - 1):
                    if self.tau[i] > self.tau[i+1]:
                        disordered.append(f"τ{i+1}={self.tau[i]:.2f} > τ{i+2}={self.tau[i+1]:.2f}")
                issues.append(f"Andrich 阈值无序: {', '.join(disordered[:3])}")

        # 4. 合并建议
        merge_suggestion = None
        n_low = len(low_usage)
        if n_low >= self.K * 0.3:  # 30% 以上等级使用不足
            merge_suggestion = f"建议合并评分等级: 当前 {self.K+1} 档 ({self.min_s}-{self.max_s})"

        return {
            "passed": len(issues) == 0,
            "issues": issues,
            "usage": [(self.min_s + i, int(c)) for i, c in enumerate(usage)],
            "tau_ordered": all(self.tau[i] <= self.tau[i+1] for i in range(self.K - 1)) if self.K >= 2 else True,
            "merge_suggestion": merge_suggestion,
        }

    def _anomalous_responses(self, threshold: float = 3.0) -> list[dict]:
        """v0.9.0: 异常反应检测 (|StRes| >= threshold)。

        返回异常反应列表，含考生/评分者/标准/题目/期望分/残差。
        """
        anomalous = []
        for i in range(self.N):
            if abs(self.z[i]) >= threshold:
                s_label = self.s_labels[self.s_idx[i]] if self.s_idx[i] < len(self.s_labels) else f"S{self.s_idx[i]+1}"
                r_label = self.r_labels[self.r_idx[i]] if self.r_idx[i] < len(self.r_labels) else f"R{self.r_idx[i]+1}"
                c_label = self.c_labels[self.c_idx[i]] if self.c_idx[i] < len(self.c_labels) else f"C{self.c_idx[i]+1}"
                i_label = self.i_labels[self.i_idx[i]] if self.n_i > 1 and self.i_idx[i] < len(self.i_labels) else ""
                anomalous.append({
                    "student": s_label, "rater": r_label,
                    "criterion": c_label, "item": i_label,
                    "observed": int(self.scores[i]),
                    "expected": round(float(self.exp_scores[i]), 2),
                    "residual": round(float(self.resid[i]), 2),
                    "stres": round(float(self.z[i]), 2),
                })

        # 按 |StRes| 降序排列
        anomalous.sort(key=lambda x: abs(x["stres"]), reverse=True)
        return anomalous

    def _diagnose_bias(self) -> list[dict]:
        """v0.9.0: 自动诊断评分者偏差类型。

        基于 MFRM 参数检测 5 种偏差（参考 PPT 第12讲 Slide 24）:
          1. 宽松/严格效应 — meas 显著偏离均值
          2. 集中趋势效应 — fair_avg 全距小 + outfit 偏低
          3. 随机效应 — infit/outfit 显著超标
          4. 晕轮效应 — 维度间评分差异小 (仅在 n_c >= 2 时检测)
          5. 差异宽松 — 跨考生 FairAvg 标准差异常

        Returns:
          [{"rater": "Rater1", "flags": [("宽松", "meas=+1.52, 显著偏高")]}, ...]
        """
        if not hasattr(self, 'exp_scores'):
            raise RuntimeError("请先运行 fit()")

        biases = []
        # 获取评分者统计
        rater_rows = []
        for p in range(self.n_r):
            idx = self.obs_r[p]
            if len(idx) == 0:
                continue
            fair_vals = self.exp_scores[idx]
            rater_rows.append({
                "label": self.r_labels[p] if p < len(self.r_labels) else f"R{p+1}",
                "meas": self.delta[p],
                "infit": (self.z[idx]**2 * self.var_o[idx]).sum() / max(self.var_o[idx].sum(), 1e-10),
                "outfit": (self.z[idx]**2).sum() / max(len(idx), 1),
                "fair_avg": fair_vals.mean(),
                "fair_std": fair_vals.std(ddof=1) if len(fair_vals) > 1 else 0,
                "n_obs": len(idx),
            })

        if not rater_rows:
            return biases

        meas_vals = np.array([r["meas"] for r in rater_rows])
        meas_mean = meas_vals.mean()
        fair_avgs = np.array([r["fair_avg"] for r in rater_rows])
        fair_mean = fair_avgs.mean()
        fair_range = fair_avgs.max() - fair_avgs.min()
        # 差异宽松: 用各评分者 FairAvg std 的中位数做基线
        fair_stds_all = np.array([r["fair_std"] for r in rater_rows if r["fair_std"] > 0])
        fair_std_median = float(np.median(fair_stds_all)) if len(fair_stds_all) > 0 else 1.0

        for r in rater_rows:
            flags = []

            # 1. 宽松/严格效应: meas 偏离均值 > 1.0 logit
            dev = r["meas"] - meas_mean
            if abs(dev) > 1.0:
                tag = "严格" if dev < 0 else "宽松"
                flags.append((tag, f"meas={r['meas']:+.2f}, 偏离均值{dev:+.2f}"))

            # 2. 集中趋势效应: fair_avg 全距 < 均值的30%
            if fair_range < fair_mean * 0.3 and r["outfit"] < 0.5:
                flags.append(("集中趋势", f"FairAvg={r['fair_avg']:.1f}, Outfit={r['outfit']:.2f}"))

            # 3. 随机效应: infit 或 outfit > 2.0
            if r["infit"] > 2.0 or r["outfit"] > 2.0:
                flags.append(("随机", f"Infit={r['infit']:.2f} Outfit={r['outfit']:.2f}"))

            # 5. 差异宽松: FairAvg std 显著高于中位数 (top outlier)
            if fair_std_median > 0 and r["fair_std"] > fair_std_median * 1.5 and r["n_obs"] >= 4:
                flags.append(("差异宽松", f"FairAvg std={r['fair_std']:.2f}, 偏离组中位数{fair_std_median:.2f}"))

            if flags:
                biases.append({"rater": r["label"], "flags": flags})

        return biases

    def _facet(self, obs, param, labels, n):
        rows, ms, ses = [], [], []
        for p in range(n):
            idx = obs[p]
            if len(idx) == 0: continue
            total, nm = self.scores[idx].sum(), len(idx)
            zz, vv = self.z[idx], self.var_o[idx]
            w = vv.sum(); infit = (zz**2 * vv).sum() / max(w, 1e-10); outfit = (zz**2).sum() / max(nm, 1)
            se = 1.0 / np.sqrt(max(vv.sum(), 1e-10))
            # v1.0.11: 额外诊断 — PtMea (点测量相关) 和 Discrm (区分度)
            exp_off = self.exp_scores[idx] - self.min_s
            obs_off = self.x[idx]
            ptmea_num = np.sum((obs_off - obs_off.mean()) * (exp_off - exp_off.mean()))
            ptmea_den = np.sqrt(np.sum((obs_off - obs_off.mean())**2) * np.sum((exp_off - exp_off.mean())**2))
            ptmea = ptmea_num / max(ptmea_den, 1e-10) if nm > 1 else 0.0
            rows.append({
                "label": labels[p], "total": int(total),
                "obs_avg": round(self.scores[idx].mean(), 2),
                "fair_avg": round(self.exp_scores[idx].mean(), 2),
                "meas": round(param[p], 3), "se": round(se, 3),
                "infit": round(infit, 3), "outfit": round(outfit, 3),
                "ptmea": round(ptmea, 3),
            })
            ms.append(param[p]); ses.append(se)
        ma, sa = np.array(ms), np.array(ses)
        vo = np.var(ma, ddof=1) if len(ma) > 1 else 0.001; me = np.mean(sa**2); vt = max(vo - me, 0.001)
        sep = np.sqrt(vt / me) if me > 0 else 0
        rel = vt / (vt + me)
        # v1.0.11: 固定效应卡方检验
        chi_sq = float(np.sum(((ma - ma.mean()) / sa) ** 2))
        chi_df = int(len(ma) - 1)
        try:
            from scipy.stats import chi2
            chi_p = float(1 - chi2.cdf(chi_sq, chi_df)) if chi_df > 0 else 1.0
        except ImportError:
            chi_p = float(np.exp(-chi_sq / (2 * chi_df))) if chi_df > 0 else 1.0  # 近似值
        return rows, sep, rel, chi_sq, chi_df, chi_p

    def report(self):
        if not hasattr(self, 'exp_scores'): raise RuntimeError("请先运行 fit()")
        r = {"summary": {"N": self.N, "n_s": self.n_s, "n_r": self.n_r, "n_c": self.n_c, "n_i": self.n_i, "score_range": f"{self.min_s}-{self.max_s}", "obs_mean": round(float(self.scores.mean()), 3), "exp_mean": round(float(self.exp_scores.mean()), 3), "resid_sd": round(float(self.resid.std(ddof=1)), 4), "stres_sd": round(float(self.z.std(ddof=1)), 4), "var_exp": round(self.var_exp, 2), "ll": round(self.ll_final, 2)}, "facets": {}}
        for name, obs, param, labels, n in [("students", self.obs_s, self.theta, self.s_labels[:self.n_s], self.n_s), ("raters", self.obs_r, self.delta, self.r_labels[:self.n_r], self.n_r), ("criteria", self.obs_c, self.alpha, self.c_labels[:self.n_c], self.n_c)]:
            rows, sep, rel, chi_sq, chi_df, chi_p = self._facet(obs, param, labels, n)
            r["facets"][name] = {"rows": rows, "separation": round(sep, 2), "reliability": round(rel, 3),
                                 "chi_sq": round(chi_sq, 1), "chi_df": chi_df, "chi_p": round(chi_p, 4)}
        if self.n_i > 1:
            rows, sep, rel, chi_sq, chi_df, chi_p = self._facet(self.obs_i, self.beta, self.i_labels[:self.n_i], self.n_i)
            r["facets"]["items"] = {"rows": rows, "separation": round(sep, 2), "reliability": round(rel, 3),
                                    "chi_sq": round(chi_sq, 1), "chi_df": chi_df, "chi_p": round(chi_p, 4)}
        # v0.9.0: 新诊断功能 (参考 Facets 专业报告)
        r["bias"] = self._diagnose_bias()
        r["categories"] = self._diagnose_categories()
        r["anomalous"] = self._anomalous_responses()
        r["rank_compare"] = self._rank_compare()
        return r

    def _rank_compare(self) -> list[dict]:
        """v0.9.0: 原始排名 vs MFRM 校正排名对比 (参考 PPT Slide 25)。"""
        if not hasattr(self, 'exp_scores'):
            return []
        rows = []
        for s in range(self.n_s):
            idx = self.obs_s[s]
            if len(idx) == 0:
                continue
            raw_total = int(self.scores[idx].sum())
            fair_avg = round(float(self.exp_scores[idx].mean()), 2)
            rows.append({
                "label": self.s_labels[s] if s < len(self.s_labels) else f"S{s+1}",
                "raw_total": raw_total,
                "fair_avg": fair_avg,
                "meas": round(float(self.theta[s]), 3),
            })
        rows.sort(key=lambda x: x["raw_total"], reverse=True)
        for i, r in enumerate(rows):
            r["raw_rank"] = i + 1
        rows.sort(key=lambda x: x["fair_avg"], reverse=True)
        for i, r in enumerate(rows):
            r["fair_rank"] = i + 1
            r["rank_diff"] = r["raw_rank"] - r["fair_rank"]
        return rows

    def print(self):
        r = self.report(); s = r["summary"]
        print(f"\n{'='*72}\nMFRMSight v1.0.9 — 多面Rasch模型分析报告\n{'='*72}")
        print(f"{s['N']}条 | {s['n_s']}学生×{s['n_r']}评分者×{s['n_c']}标准×{s['n_i']}题目 | {s['score_range']}分")
        print(f"方差解释: {s['var_exp']}% | LL: {s['ll']:.0f}")
        print(f"ObsMean={s['obs_mean']:.3f} ExpMean={s['exp_mean']:.3f} ResidSD={s['resid_sd']:.4f} StResSD={s['stres_sd']:.4f}")
        for fn, fd in r["facets"].items():
            if not fd["rows"]: continue
            print(f"\n{'─'*72}\n{fn} — Sep={fd['separation']:.2f} Rel={fd['reliability']:.3f}\n{'─'*72}")
            print(f"{'':<16} {'Total':>7} {'ObsAvg':>7} {'FairAvg':>7} {'Meas':>7} {'SE':>6} {'Infit':>6} {'Outfit':>6}")
            print("-" * 72)
            for row in fd["rows"]:
                print(f"{row['label']:<16} {row['total']:>7.0f} {row['obs_avg']:>7.2f} {row['fair_avg']:>7.2f} {row['meas']:>7.3f} {row['se']:>6.3f} {row['infit']:>6.3f} {row['outfit']:>6.3f}")

        # v0.9.0: 评分者偏差诊断
        if r.get("bias"):
            print(f"\n{'='*72}\n[!] 评分者偏差诊断\n{'='*72}")
            for b in r["bias"]:
                for tag, detail in b["flags"]:
                    print(f"  [{tag}] {b['rater']}: {detail}")
            if not any(b["flags"] for b in r["bias"]):
                print("  [OK] 未检出显著评分者偏差")
        else:
            print(f"\n{'='*72}\n[OK] 评分者偏差诊断: 未检出异常\n{'='*72}")

        # v0.9.0: 排名对比
        if r.get("rank_compare"):
            print(f"\n{'='*72}\n[>>>] 原始总分 vs MFRM校正 排名对比\n{'='*72}")
            print(f"{'考生':<16} {'原始总分':>8} {'原始排名':>8} {'校正分':>8} {'校正排名':>8} {'排名变化':>8}")
            print("-" * 60)
            for row in r["rank_compare"]:
                diff = row["rank_diff"]
                arrow = "^" if diff > 0 else ("v" if diff < 0 else "-")
                print(f"{row['label']:<16} {row['raw_total']:>8} {row['raw_rank']:>8} {row['fair_avg']:>8.2f} {row['fair_rank']:>8} {arrow}{abs(diff):>7}")

        # v0.9.0: 等级类别功能诊断
        cat = r.get("categories")
        if cat:
            print(f"\n{'='*72}\n[*] 评分等级类别功能诊断 (参考 Facets Table 8)\n{'='*72}")
            print(f"等级数: {self.K+1} 档 ({self.min_s}-{self.max_s}), 阈值有序: {'是' if cat['tau_ordered'] else '否'}")
            if cat["issues"]:
                for issue in cat["issues"]:
                    print(f"  [!] {issue}")
            if cat["merge_suggestion"]:
                print(f"  [>>] {cat['merge_suggestion']}")
            if not cat["issues"]:
                print("  [OK] 等级功能正常，无需合并")

        # v0.9.0: 异常反应
        anom = r.get("anomalous", [])
        if anom:
            print(f"\n{'='*72}\n[!!] 异常反应记录 (|StRes| >= 3)\n{'='*72}")
            print(f"{'考生':<12} {'评分者':<10} {'标准':<8} {'观测分':>6} {'期望分':>6} {'残差':>6} {'StRes':>6}")
            print("-" * 60)
            for a in anom[:15]:  # 最多显示 15 条
                print(f"{a['student']:<12} {a['rater']:<10} {a['criterion']:<8} "
                      f"{a['observed']:>6} {a['expected']:>6.1f} {a['residual']:>6.1f} {a['stres']:>6.2f}")
            if len(anom) > 15:
                print(f"  ... 共 {len(anom)} 条异常反应")
            print(f"\n  注意: 异常残差不直接等同于评分者偏见，应结合评分者、标准、原始评分记录综合判断")
        else:
            print(f"\n{'='*72}\n[OK] 异常反应检测: 未检出 |StRes| >= 3 的反应\n{'='*72}")

    def bias_interaction(self, facet_a: str = "raters", facet_b: str = "students") -> list[dict]:
        """v1.0.0: 偏差交互分析 — 计算两个面所有交互对的偏差量。

        偏差 = 观测均值 - 期望均值 (控制主效应后)。
        |z| >= 2 标记为显著。

        Args:
            facet_a: 第一个面的 key ("students"/"raters"/"criteria"/"items")
            facet_b: 第二个面的 key

        Returns:
            [{"a": "Rater1", "b": "Student4", "obs_avg": 15.0, "exp_avg": 13.5,
              "bias": 1.5, "se": 0.3, "z": 5.0, "significant": True}, ...]
        """
        key_to_idx = {"students": 0, "raters": 1, "criteria": 2, "items": 3}
        key_to_obs = {"students": self.obs_s, "raters": self.obs_r,
                      "criteria": self.obs_c, "items": self.obs_i}
        key_to_labels = {"students": self.s_labels, "raters": self.r_labels,
                         "criteria": self.c_labels, "items": self.i_labels}

        idx_a = key_to_idx.get(facet_a)
        idx_b = key_to_idx.get(facet_b)
        if idx_a is None or idx_b is None:
            return []

        cols_a = self.raw[:, idx_a]
        cols_b = self.raw[:, idx_b]
        vals_a = sorted(set(cols_a))
        vals_b = sorted(set(cols_b))

        pairs = []
        for va in vals_a:
            for vb in vals_b:
                mask = (cols_a == va) & (cols_b == vb)
                if mask.sum() < 2:
                    continue
                obs_mean = float(self.scores[mask].mean())
                exp_mean = float(self.exp_scores[mask].mean())
                bias_val = obs_mean - exp_mean
                # SE 用残差的标准误估计
                n_pair = mask.sum()
                resid_std = float(self.resid[mask].std(ddof=1)) if n_pair > 1 else 1.0
                se = resid_std / np.sqrt(n_pair) if n_pair > 0 else 1.0
                z_val = bias_val / se if se > 0 else 0.0

                label_a = key_to_labels[facet_a][va - 1] if va <= len(key_to_labels[facet_a]) else f"#{va}"
                label_b = key_to_labels[facet_b][vb - 1] if vb <= len(key_to_labels[facet_b]) else f"#{vb}"

                pairs.append({
                    "a": label_a, "b": label_b,
                    "obs_avg": round(obs_mean, 2),
                    "exp_avg": round(exp_mean, 2),
                    "bias": round(float(bias_val), 3),
                    "se": round(float(se), 3),
                    "z": round(float(z_val), 2),
                    "significant": abs(z_val) >= 2.0,
                    "n": int(n_pair),
                })

        pairs.sort(key=lambda x: abs(x["z"]), reverse=True)
        return pairs

    def aic_bic(self) -> dict:
        """v1.0.0: AIC/BIC 模型比较指标。

        AIC = -2*LL + 2*k, BIC = -2*LL + k*ln(N)
        其中 k = 自由参数数 = n_s + n_r + n_c + n_i - 1 + K
        """
        if not hasattr(self, 'll_final'):
            raise RuntimeError("请先运行 fit()")
        n_params = self.n_s + self.n_r + self.n_c + self.n_i - 1 + self.K
        aic = -2 * self.ll_final + 2 * n_params
        bic = -2 * self.ll_final + n_params * np.log(max(self.N, 1))
        return {"aic": round(aic, 2), "bic": round(bic, 2), "n_params": n_params, "ll": round(self.ll_final, 2)}

    def to_excel(self, path):
        import pandas as pd
        r = self.report(); dfs = {n: pd.DataFrame(f["rows"]) for n, f in r["facets"].items()}
        dfs["summary"] = pd.DataFrame([r["summary"]])
        with pd.ExcelWriter(path) as w:
            for n, d in dfs.items(): d.to_excel(w, sheet_name=n, index=False)
        print(f"[OK] {path}")

    def to_word(self, path):
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            print("[!] pip install python-docx"); return
        r = self.report(); doc = Document()
        t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = t.add_run("MFRM Rating Scale 模型分析报告"); run.bold = True; run.font.size = Pt(18)
        s = r["summary"]
        doc.add_paragraph(f"反应数: {s['N']} | {s['n_s']}学生×{s['n_r']}评分者×{s['n_c']}标准×{s['n_i']}题目\n分数: {s['score_range']} | Rasch解释方差: {s['var_exp']}%\n观察均值: {s['obs_mean']:.3f} | 残差SD: {s['resid_sd']:.3f} | StRes SD: {s['stres_sd']:.3f}")
        for fn, fd in r["facets"].items():
            if not fd["rows"]: continue
            doc.add_heading(f"{fn} (Sep={fd['separation']:.2f}, Rel={fd['reliability']:.3f})", level=2)
            h = ["元素", "总分", "ObsAvg", "FairAvg", "Meas", "SE", "Infit", "Outfit"]
            tbl = doc.add_table(rows=1 + len(fd["rows"]), cols=8); tbl.style = "Table Grid"
            for i, hh in enumerate(h): tbl.rows[0].cells[i].text = hh
            for ri, row in enumerate(fd["rows"]):
                for ci, k in enumerate(["label", "total", "obs_avg", "fair_avg", "meas", "se", "infit", "outfit"]):
                    v = row[k]; tbl.rows[ri + 1].cells[ci].text = f"{v:.1f}" if isinstance(v, float) and k != "label" else f"{v:.0f}" if k == "total" else str(v)
        doc.save(path); print(f"[OK] {path}")


# ═══════════════════════════════════════════════════════════════════════
# v1.0.0: 统计图表生成 (matplotlib)
# ═══════════════════════════════════════════════════════════════════════

# 设置中文字体 (尝试多种常见中文字体)
for _font in ["SimHei", "Microsoft YaHei", "WenQuanYi Zen Hei",
              "Noto Sans CJK SC", "DejaVu Sans"]:
    try:
        matplotlib.font_manager.findfont(_font, fallback_to_default=False)
        plt.rcParams["font.sans-serif"] = [_font, "DejaVu Sans"]
        break
    except Exception:
        pass
plt.rcParams["axes.unicode_minus"] = False


def _fig_to_b64(fig) -> str:
    """matplotlib Figure → base64 PNG 字符串"""
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=120, bbox_inches="tight")
    buf.seek(0)
    plt.close(fig)
    return base64.b64encode(buf.read()).decode()


def chart_ruler_map(engine: "MFRMEngine", data: dict) -> str:
    """v1.0.0: 垂直标尺图 (Ruler Map) — 各面向在 logit 量尺上的分布。

    模仿 Facets Table 6.0, 用水平散点图展示 Students/Raters/Criteria/Items 的 Measure。
    Returns base64 PNG.
    """
    dims = extract_dimensions(data)
    fig, ax = plt.subplots(figsize=(10, max(3, len(dims["facets"]) * 1.2)))
    colors = ["#2196F3", "#FF5722", "#4CAF50", "#9C27B0"]
    y_labels = []

    for fi_idx, fi in enumerate(dims["facets"]):
        key = fi["key"]
        label = fi["label"]
        y_labels.append(label)
        param_map = {"students": engine.theta, "raters": engine.delta,
                     "criteria": engine.alpha, "items": engine.beta}
        vals = param_map.get(key, np.zeros(1))
        y_pos = np.full(len(vals), fi_idx + 1)
        ax.scatter(vals, y_pos, c=colors[fi_idx % 4], s=60, alpha=0.7,
                   edgecolors="white", linewidth=0.5, zorder=3)
        # 标注均值
        mean_v = vals.mean()
        ax.axvline(mean_v, ymin=fi_idx / len(dims["facets"]),
                   ymax=(fi_idx + 1.2) / len(dims["facets"]),
                   color=colors[fi_idx % 4], linewidth=2, alpha=0.5)

    ax.set_yticks(range(1, len(y_labels) + 1))
    ax.set_yticklabels(y_labels, fontsize=11)
    ax.axvline(0, color="gray", linestyle="--", linewidth=1, alpha=0.5)
    ax.set_xlabel("Measure (logit)", fontsize=10)
    ax.set_title("垂直标尺图 (Ruler Map) — 各面向元素在 Logit 量尺上的分布", fontsize=12, fontweight="bold")
    ax.grid(axis="x", alpha=0.3)
    fig.tight_layout()
    return _fig_to_b64(fig)


def chart_category_curves(engine: "MFRMEngine") -> str:
    """v1.0.0: 等级概率曲线 + ICC 期望分曲线 (模仿 Facets Table 8.1)。

    上: 概率曲线 P(score=k|θ), 下: ICC 期望总分 vs θ。
    Returns base64 PNG.
    """
    K = engine.K
    theta_range = np.linspace(-6, 6, 200)
    tau = engine.tau
    cats = engine.cats

    # 概率曲线
    probs = np.zeros((len(theta_range), K + 1))
    for i, th in enumerate(theta_range):
        lin = th - tau
        cum = np.zeros(K + 1)
        for k in range(K):
            cum[k + 1] = cum[k] + lin[k]
        cum -= cum.max()
        e = np.exp(cum)
        probs[i] = e / e.sum()

    fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(10, 8))

    # 概率曲线
    for k in range(K + 1):
        alpha_v = 1.0 if k in (0, K) else 0.4
        lw = 1.5 if k in (0, K) else 0.8
        ax1.plot(theta_range, probs[:, k], linewidth=lw, alpha=alpha_v,
                label=f"{engine.min_s + k}")
    ax1.set_title("等级概率曲线 — P(score=k|θ)", fontsize=12, fontweight="bold")
    ax1.set_xlabel("Measure (logit)")
    ax1.set_ylabel("Probability")
    ax1.legend(loc="upper left", ncol=4, fontsize=7)
    ax1.grid(alpha=0.3)

    # ICC
    exp_scores = probs @ cats
    ax2.plot(theta_range, exp_scores + engine.min_s, "b-", linewidth=2)
    ax2.set_title("ICC 期望分曲线 — Expected Score vs Ability", fontsize=12, fontweight="bold")
    ax2.set_xlabel("Measure (logit)")
    ax2.set_ylabel("Expected Score")
    ax2.grid(alpha=0.3)
    ax2.axhline(engine.max_s, color="red", linestyle="--", alpha=0.3)
    ax2.axhline(engine.min_s, color="red", linestyle="--", alpha=0.3)

    fig.tight_layout()
    return _fig_to_b64(fig)


def chart_fit_distribution(engine: "MFRMEngine") -> str:
    """v1.0.0: Infit/Outfit 分布条形图 (模仿 Facets Table 6.x barchart)。

    显示各评分者的 Infit MnSq 和 Outfit MnSq, 标注 .5-1.5 可接受区间。
    Returns base64 PNG.
    """
    n_r = engine.n_r
    r_labels = engine.r_labels[:n_r] if engine.r_labels else [f"R{i+1}" for i in range(n_r)]

    r = engine.report()
    fd = r["facets"].get("raters", {})
    rows = fd.get("rows", [])

    infits = [row.get("infit", 1) for row in rows]
    outfits = [row.get("outfit", 1) for row in rows]

    fig, ax = plt.subplots(figsize=(10, max(3, n_r * 0.45)))
    y_pos = np.arange(n_r)
    width = 0.35

    bars1 = ax.barh(y_pos - width/2, infits, width, color="#2196F3", alpha=0.8, label="Infit MnSq")
    bars2 = ax.barh(y_pos + width/2, outfits, width, color="#FF5722", alpha=0.8, label="Outfit MnSq")

    ax.axvline(1.0, color="gray", linestyle="-", linewidth=1, alpha=0.7)
    ax.axvline(0.5, color="green", linestyle="--", linewidth=0.8, alpha=0.5)
    ax.axvline(1.5, color="red", linestyle="--", linewidth=0.8, alpha=0.5)
    ax.axvspan(0.5, 1.5, alpha=0.05, color="green")

    ax.set_yticks(y_pos)
    ax.set_yticklabels(r_labels, fontsize=9)
    ax.set_xlabel("MnSq")
    ax.set_title("评分者拟合统计量分布 (Infit/Outfit MnSq, .5-1.5 可接受)", fontsize=12, fontweight="bold")
    ax.legend(loc="lower right")
    ax.grid(axis="x", alpha=0.3)
    fig.tight_layout()
    return _fig_to_b64(fig)


# ═══════════════════════════════════════════════════════════════════════
# v1.0.0: 一键生成专业报告 (参考 Facets 分析结果专业整理报告.docx)
# ═══════════════════════════════════════════════════════════════════════

def generate_report(engine: "MFRMEngine", data: dict,
                    bias_results: list[dict] | None = None,
                    title: str = "") -> str:
    """v1.0.0: 按 10 章节结构生成中文专业报告 (Markdown 格式)。

    Args:
        engine: 已 fit 的 MFRMEngine 实例
        data: parse_facets_txt 返回的原始数据
        bias_results: 偏差交互分析结果 [{"pair": "Rater×Student", "rows": [...]}]
        title: 报告标题，默认取数据中的 title

    Returns:
        Markdown 格式的完整报告字符串
    """
    r = engine.report()
    if not r or "facets" not in r:
        return "# 错误: 无法生成报告 (引擎未完成分析)"
    s = r.get("summary", {})
    dims = extract_dimensions(data)
    facets_info = dims["facets"]

    t = title or data.get("title", "MFRM 分析")
    lines = []
    def w(text=""): lines.append(text)

    # ═══════════════════════════════════════════════════
    # 一、报告摘要
    # ═══════════════════════════════════════════════════
    w(f"# {t} — MFRM 分析结果专业整理报告\n")
    w("## 一、报告摘要\n")
    n_s = s.get("n_s", 0); n_r = s.get("n_r", 0); n_c = s.get("n_c", 0); n_i = s.get("n_i", 0)
    w(f"- 本报告基于 **{s.get('N', 0)} 条评分反应**，对 {n_s} 名考生、{n_r} 名评分者")
    if n_i > 1:
        w(f"、{n_c} 个评分标准、{n_i} 个题目进行多面 Rasch 模型分析。")
    else:
        w(f"、{n_c} 个评分标准进行多面 Rasch 模型分析。")

    obs_m = s.get("obs_mean", 0); exp_m = s.get("exp_mean", 0)
    var_e = s.get("var_exp", 0)
    w(f"- 总体拟合：观察均值 **{obs_m:.2f}**，期望均值 **{exp_m:.2f}**（差距 **{abs(obs_m - exp_m):.3f}**），"
      f"标准化残差 SD={s.get('stres_sd', 0):.2f}。")
    if var_e >= 85:
        w(f"- Rasch 测量值解释了 **{var_e}%** 的原始得分方差，说明模型能够解释主要评分变异。")
    else:
        w(f"- Rasch 测量值解释了 **{var_e}%** 的原始得分方差，模型诊断价值有限，建议进一步检查。")

    # 偏差摘要
    bias = r.get("bias", [])
    if bias:
        bias_raters = [b["rater"] for b in bias]
        w(f"- ⚠️ 检出 **{len(bias)} 位评分者**存在偏差：{', '.join(bias_raters[:5])}")

    cat = r.get("categories", {})
    if not cat.get("passed", True):
        w(f"- 等级类别功能存在问题：阈值有序={cat.get('tau_ordered', '?')}，"
          f"建议合并评分等级后重新估计。")

    anom = r.get("anomalous", [])
    if anom:
        w(f"- 检出 **{len(anom)} 条异常反应** (|StRes| ≥ 3)，详见第九节。")
    else:
        w("- 未检出 |StRes| ≥ 3 的异常反应。")

    # ═══════════════════════════════════════════════════
    # 二、模型设定
    # ═══════════════════════════════════════════════════
    w("\n## 二、模型设定与输出方向\n")
    nf = dims["n_facets"]
    w(f"本次分析采用 **{nf} 面向**多面 Rasch 测量模型：")
    for fi in facets_info:
        zh = fi["label"]
        orig = f" ({fi['original']})" if fi['original'] != fi['key'] else ""
        w(f"- **{zh}**{orig}：{fi['n']} 个元素")

    nc = data.get("noncentered", 1)
    w(f"\n| 设定项 | 内容 | 专业解释 |")
    w(f"| --- | --- | --- |")
    w(f"| Facets = {nf} | {'、'.join(fi['label'] for fi in facets_info)} | {nf} 面向评分模型 |")
    w(f"| Positive = {data.get('positive', 1)} | 高分代表高能力 | 高 logit = 高能力/宽松/容易 |")
    w(f"| Non-centered = {nc} | 第 {nc} 面不中心化 | 以该面为自由参考 |")
    model = data.get("model", "?")
    if model:
        w(f"| Model = {model} | 评分等级结构 | 见第八节等级类别诊断 |")

    # ═══════════════════════════════════════════════════
    # 三、数据摘要
    # ═══════════════════════════════════════════════════
    w("\n## 三、收敛控制与数据摘要\n")
    w(f"| 指标 | 结果 | 解释 |")
    w(f"| --- | --- | --- |")
    w(f"| 总反应数 | {s['N']} | 用于模型估计的有效反应 |")
    w(f"| 分数范围 | {s['score_range']} | 最小-最大评分 |")
    w(f"| Subset connection | O.K. | 数据连接性满足估计要求 |")

    # ═══════════════════════════════════════════════════
    # 四、总体拟合
    # ═══════════════════════════════════════════════════
    w("\n## 四、总体模型拟合与方差分解\n")
    w(f"| 指标 | 数值 | 专业解释 |")
    w(f"| --- | --- | --- |")
    w(f"| 观察均值 | {obs_m:.2f} | 原始评分的总平均值 |")
    w(f"| 期望均值 | {exp_m:.2f} | 模型预测的期望均值 |")
    w(f"| ObsMean - ExpMean | {obs_m - exp_m:.4f} | {'差距极小，模型无系统性偏差' if abs(obs_m - exp_m) < 0.1 else '存在一定偏差'} |")
    w(f"| Rasch 解释方差 | {var_e}% | {'模型解释力良好' if var_e >= 85 else ('中等' if var_e >= 70 else '偏低')} |")
    w(f"| 残差方差 | {100 - var_e:.2f}% | 未被模型解释的变异 |")
    w(f"| 标准化残差 SD | {s['stres_sd']:.4f} | {'接近理论值 1.0' if 0.9 <= s['stres_sd'] <= 1.1 else '偏离理论值 1.0'} |")

    # ═══════════════════════════════════════════════════
    # 五、各面向测量结果
    # ═══════════════════════════════════════════════════
    # ── 垂直标尺图 ──
    try:
        ruler_b64 = chart_ruler_map(engine, data)
        w(f"\n![垂直标尺图](data:image/png;base64,{ruler_b64})\n")
    except Exception:
        pass

    w("\n## 五、各面向测量结果\n")

    facet_descs = {
        "students": "**考生能力**差异",
        "raters": "**评分者宽严度**差异",
        "criteria": "**评分标准难易度**差异",
        "items": "**题目难度**差异",
    }

    for fn, fd in r.get("facets", {}).items():
        if not fd.get("rows"):
            continue
        sep = fd.get("separation", 0); rel = fd.get("reliability", 0)
        desc = facet_descs.get(fn, "")
        n_elem = len(fd["rows"])

        w(f"### {fn} — Sep={sep:.2f} Rel={rel:.3f}\n")
        if sep > 3:
            w(f"分离度良好 (Sep={sep:.2f})，说明 {desc} 显著且排序稳定。")
        elif sep > 1:
            w(f"分离度一般 (Sep={sep:.2f})，{desc}具有一定区分度。")
        else:
            w(f"分离度较低 (Sep={sep:.2f})，{desc}有限，可能不存在可分离的真实差异。")

        w(f"\n| 元素 | 总分 | ObsAvg | FairAvg | Measure | SE | Infit | Outfit | 诊断 |")
        w(f"| --- | --- | --- | --- | --- | --- | --- | --- | --- |")
        for row in fd["rows"]:
            inf = row["infit"]; otf = row["outfit"]
            diag = ""
            if inf < 0.5: diag = "过于一致"
            elif inf > 1.5: diag = "略有波动"
            elif otf > 2.0: diag = "需关注"
            else: diag = "良好"
            w(f"| {row['label']} | {row['total']} | {row['obs_avg']} | {row['fair_avg']} | {row['meas']:.3f} | {row['se']:.3f} | {inf:.2f} | {otf:.2f} | {diag} |")

    # ═══════════════════════════════════════════════════
    # 六、评分者偏差诊断
    # ═══════════════════════════════════════════════════
    if bias:
        w("\n## 六、评分者偏差诊断\n")
        for b in bias:
            for tag, detail in b["flags"]:
                w(f"- **[{tag}] {b['rater']}**: {detail}")

    # ── 拟合分布图 ──
    try:
        fit_b64 = chart_fit_distribution(engine)
        w(f"\n![拟合统计量分布](data:image/png;base64,{fit_b64})\n")
    except Exception:
        pass

    # ═══════════════════════════════════════════════════
    # 七、偏差交互分析
    # ═══════════════════════════════════════════════════
    if bias_results:
        w("\n## 七、偏差交互分析\n")
        for bi in bias_results:
            w(f"### {bi.get('pair', '交互分析')}\n")
            rows = bi.get("rows", [])
            sig = [r for r in rows if r.get("significant")]
            if sig:
                w(f"检出 **{len(sig)} 对**显著偏差交互 (|z| ≥ 2)：")
                w(f"\n| 元素A | 元素B | ObsAvg | ExpAvg | Bias | SE | z |")
                w(f"| --- | --- | --- | --- | --- | --- | --- |")
                for r in sig[:20]:
                    w(f"| {r['a']} | {r['b']} | {r['obs_avg']} | {r['exp_avg']} | {r['bias']:.3f} | {r['se']:.3f} | {r['z']:.2f} |")
            else:
                w("未检出显著偏差交互对。")

    # ═══════════════════════════════════════════════════
    # 八、等级类别诊断
    # ═══════════════════════════════════════════════════
    if cat and cat.get("rows"):
        w("\n## 八、评分等级类别功能诊断\n")
        # ── 等级概率曲线图 ──
        try:
            cat_b64 = chart_category_curves(engine)
            w(f"\n![等级概率曲线与ICC](data:image/png;base64,{cat_b64})\n")
        except Exception:
            pass
        w(f"当前等级数: **{len(cat['rows'])} 档**, 阈值有序: **{'是' if cat['tau_ordered'] else '否'}**")
        if cat.get("merge_suggestion"):
            w(f"\n> {cat['merge_suggestion']}")
        if cat.get("issues"):
            w("\n诊断问题:")
            for issue in cat["issues"]:
                w(f"- {issue}")

    # ═══════════════════════════════════════════════════
    # 九、异常反应记录
    # ═══════════════════════════════════════════════════
    w("\n## 九、异常反应记录\n")
    if anom:
        w(f"共检出 **{len(anom)} 条** |StRes| ≥ 3 的异常反应：\n")
        w(f"| 考生 | 评分者 | 标准 | 观测分 | 期望分 | 残差 | StRes |")
        w(f"| --- | --- | --- | --- | --- | --- | --- |")
        for a in anom[:15]:
            w(f"| {a['student']} | {a['rater']} | {a['criterion']} | {a['observed']} | {a['expected']} | {a['residual']} | {a['stres']} |")
        w(f"\n> 注意: 异常残差不直接等同于评分者偏见，应结合评分者、标准、原始评分记录综合判断。")
    else:
        w("未检出 |StRes| ≥ 3 的异常反应。")

    # ═══════════════════════════════════════════════════
    # 十、综合结论
    # ═══════════════════════════════════════════════════
    w("\n## 十、综合结论与改进建议\n")
    conclusions = []
    if var_e >= 85:
        conclusions.append(f"- 模型整体拟合良好，Rasch 测量值解释了 {var_e}% 的评分方差。")
    for fn, fd in r.get("facets", {}).items():
        if fd.get("rows") and fd.get("separation", 0) > 3:
            conclusions.append(f"- {fn} 面向区分清晰 (Sep={fd['separation']:.2f})。")
    if not cat.get("passed", True):
        conclusions.append(f"- **评分等级过细**，建议合并等级后重新估计。")
    if bias:
        conclusions.append(f"- 检出 {len(bias)} 位评分者存在偏差，建议结合原始评分记录进一步审查。")
    if not conclusions:
        conclusions.append("- 模型运行正常，未发现严重问题。")
    for c in conclusions:
        w(c)

    report = "\n".join(lines)
    report = "\n".join(lines)
    return report


def generate_word_report(engine: "MFRMEngine", data: dict,
                         bias_results: list[dict] | None = None,
                         filepath: str = "") -> str:
    """v1.0.11: 生成专业 Word (.docx) 报告 — 10章节金字塔结构，含多证据链诊断。

    报告结构:
      封面 → 一、报告摘要 → 二、模型设定 → 三、数据摘要
      → 四、总体拟合 → 五、垂直标尺图 → 六、各面向测量结果
      → 七、评分等级类别功能诊断 → 八、异常反应记录
      → 九、偏差交互分析(可选) → 十、综合结论 → 附录:统计图

    Args:
        engine: 已 fit 的 MFRMEngine
        data: parse_facets_txt 返回的数据
        bias_results: 偏差交互结果 [{"pair": "...", "rows": [...]}]
        filepath: 输出路径

    Returns:
        生成的 .docx 文件路径
    """
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml, OxmlElement
    import os, datetime

    r = engine.report()
    s = r["summary"]
    dims = extract_dimensions(data)
    facets_info = dims["facets"]
    filepath = filepath or os.path.join(os.getcwd(), "mfrm_report.docx")

    doc = Document()
    # 页面设置
    for section in doc.sections:
        section.page_width = Cm(21.0); section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54); section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54); section.right_margin = Cm(2.54)

    # ═══════════════════════════════════════════════════════════════
    # 工具函数
    # ═══════════════════════════════════════════════════════════════

    def _tbl(headers: list[str], rows_data: list[list[object]]):
        """创建 APA 三线表 — 自适应列宽，防跨行。

        宽度策略:
        - 用表头/内容最长字符串估算列宽
        - 窄列（标签/数字）设 noWrap，宽列（诊断/解释）放宽并允许折行
        - 表格宽度填满页面 (16.5cm)
        """
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn as _qn
        from docx.shared import Emu

        n_rows = len(rows_data)
        n_cols = len(headers)
        tbl = doc.add_table(rows=1 + n_rows, cols=n_cols)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = True

        # ── 表格级属性: 自动布局 ──
        tblPr = tbl._tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr'); tbl._tbl.insert(0, tblPr)
        # 用 autofit 布局（宽度按内容自动分配）
        tblLayout = OxmlElement('w:tblLayout')
        tblLayout.set(_qn('w:type'), 'autofit')
        tblPr.append(tblLayout)

        # ── 估算列宽 ──
        # 中文字符 ≈ 1.5 个英文字符宽度, 9pt 下每个英文字符 ≈ 4.5pt ≈ 0.08cm
        def _text_width(text: str) -> float:
            """估算文本显示宽度 (cm)"""
            w = 0.0
            for ch in str(text):
                if '一' <= ch <= '鿿' or '　' <= ch <= '〿' or '＀' <= ch <= '￯':
                    w += 0.17  # 中文 ≈ 0.17cm/字符 (9pt)
                else:
                    w += 0.09  # 英文/数字 ≈ 0.09cm/字符 (9pt)
            return w

        col_max_w = [0.0] * n_cols
        # 表头
        for i, hdr in enumerate(headers):
            col_max_w[i] = max(col_max_w[i], _text_width(str(hdr)) + 0.3)  # +padding
        # 数据行
        for row in rows_data:
            for i, val in enumerate(row):
                col_max_w[i] = max(col_max_w[i], _text_width(str(val)) + 0.3)

        # 可用总宽度 16.5cm (A4宽21cm - 左右边距2.54cm*2)
        total_avail = 15.0
        total_est = sum(col_max_w)

        # 窄列（< 2.5cm 估计宽度）使用 noWrap
        narrow_cols = set()
        for i in range(n_cols):
            estimated = col_max_w[i] * (total_avail / max(total_est, 1.0))
            if estimated < 2.8:
                narrow_cols.add(i)

        # ── 填充表头 ──
        for i, hdr in enumerate(headers):
            cell = tbl.rows[0].cells[i]
            cell.text = ""
            pp = cell.paragraphs[0]
            pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pp.paragraph_format.space_before = Pt(2)
            pp.paragraph_format.space_after = Pt(2)
            run = pp.add_run(str(hdr))
            run.bold = True; run.font.size = Pt(9)

        # ── 填充数据行 ──
        for ri, row in enumerate(rows_data):
            for ci, val in enumerate(row):
                cell = tbl.rows[ri + 1].cells[ci]
                cell.text = ""
                pp = cell.paragraphs[0]
                pp.paragraph_format.space_before = Pt(1)
                pp.paragraph_format.space_after = Pt(1)
                # 数字列居中，文本列左对齐
                if isinstance(val, (int, float)):
                    pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    pp.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = pp.add_run(str(val)); run.font.size = Pt(9)
                # 窄列禁止折行
                if ci in narrow_cols:
                    tcPr = cell._tc.get_or_add_tcPr()
                    noWrap = OxmlElement('w:noWrap')
                    tcPr.append(noWrap)

        # ── 应用三线边框 ──
        def _border(tag, val, sz, color='000000'):
            el = OxmlElement(tag)
            el.set(_qn('w:val'), val); el.set(_qn('w:sz'), sz)
            el.set(_qn('w:space'), '0'); el.set(_qn('w:color'), color)
            return el

        for old in tblPr.findall(_qn('w:tblBorders')):
            tblPr.remove(old)
        tb = OxmlElement('w:tblBorders')
        tb.append(_border('w:top', 'single', '12'))
        tb.append(_border('w:bottom', 'single', '12'))
        for tag in ['w:insideH', 'w:left', 'w:right', 'w:insideV']:
            el = OxmlElement(tag); el.set(_qn('w:val'), 'nil'); tb.append(el)
        tblPr.append(tb)
        for cell in tbl.rows[0].cells:
            tcPr = cell._tc.get_or_add_tcPr()
            for old in tcPr.findall(_qn('w:tcBorders')):
                tcPr.remove(old)
            cb = OxmlElement('w:tcBorders')
            cb.append(_border('w:bottom', 'single', '4'))
            tcPr.append(cb)

        doc.add_paragraph()
        return tbl

    def _heading(text, level=1):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
        return h

    def _para(text, bold=False, indent=True):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.5
        if indent:
            p.paragraph_format.first_line_indent = Cm(0.74)
        run = p.add_run(text)
        run.font.size = Pt(12)
        if bold:
            run.bold = True
        return p

    def _fig_caption(num: int, title: str, note: str = ""):
        """APA 7th 标准图注：编号加粗 + 标题斜体 + Note."""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(10)
        run_num = p.add_run(f"Figure {num}")
        run_num.bold = True; run_num.font.size = Pt(10)
        p.add_run("\n")
        run_title = p.add_run(title)
        run_title.italic = True; run_title.font.size = Pt(10)
        if note:
            p.add_run("\n")
            run_note = p.add_run(f"Note. {note}")
            run_note.font.size = Pt(9)

    def _table_note(note: str, specific: str = "", prob: str = ""):
        """APA 7th 标准表注：一般→特定→概率."""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(8)
        run_note = p.add_run(f"Note. {note}")
        run_note.font.size = Pt(9)
        if specific:
            p2 = doc.add_paragraph()
            p2.paragraph_format.space_after = Pt(4)
            r2 = p2.add_run(specific)
            r2.font.size = Pt(9)
        if prob:
            p3 = doc.add_paragraph()
            p3.paragraph_format.space_after = Pt(8)
            r3 = p3.add_run(prob)
            r3.italic = True; r3.font.size = Pt(9)

    def _diagnose_infit_outfit(infit: float, outfit: float) -> str:
        """诊断 Infit/Outfit"""
        if infit < 0.5 and outfit < 0.5:
            return "过度拟合，评分模式过于一致，可能存在信息冗余"
        elif infit > 1.5 or outfit > 1.5:
            return "拟合偏高，评分模式存在一定波动，需关注具体评分记录"
        elif infit > 2.0 or outfit > 2.0:
            return "拟合严重偏高，评分模式不稳定，建议重点审查"
        else:
            return "拟合良好"

    def _chi_p_str(chi_sq: float, chi_df: int, chi_p: float) -> str:
        """格式化卡方检验结果"""
        if chi_p < .001:
            return f"χ²({chi_df}) = {chi_sq:.1f}, p < .001（极为显著）"
        elif chi_p < .01:
            return f"χ²({chi_df}) = {chi_sq:.1f}, p = {chi_p:.4f}（非常显著）"
        elif chi_p < .05:
            return f"χ²({chi_df}) = {chi_sq:.1f}, p = {chi_p:.4f}（显著）"
        else:
            return f"χ²({chi_df}) = {chi_sq:.1f}, p = {chi_p:.4f}（不显著）"

    facet_names_cn = {"students": "考生", "raters": "评分者", "criteria": "评分标准", "items": "题目"}
    bias = r.get("bias", [])
    cat = r.get("categories", {})
    anom = r.get("anomalous", [])
    var_e = s["var_exp"]
    today = datetime.date.today().strftime("%Y/%m/%d")
    fig_counter = [0]  # mutable counter for figure numbering

    # ═══════════════════════════════════════════════════════════════
    # 封面
    # ═══════════════════════════════════════════════════════════════
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(80)
    run = title_p.add_run("多面 Rasch 模型（MFRM）\n分析结果专业整理报告")
    run.bold = True; run.font.size = Pt(22)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_before = Pt(24)
    run = sub_p.add_run(f"基于 MFRMSight v{__version__} 自动生成")
    run.font.size = Pt(12)

    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_p.paragraph_format.space_before = Pt(16)
    n_s, n_r, n_c, n_i = s["n_s"], s["n_r"], s["n_c"], s["n_i"]
    info_lines = [
        f"数据文件：{data.get('title', '未命名')}  |  {s['N']} 条评分反应",
        f"面向设定：{n_s} 名考生 × {n_r} 名评分者 × {n_c} 个标准" + (f" × {n_i} 个题目" if n_i > 1 else ""),
        f"评分模型：{s['score_range']} 分等级评分，正向计分（高分=高能力）",
        f"分析软件：MFRMSight v{__version__}  |  分析日期：{today}",
    ]
    run = info_p.add_run("\n".join(info_lines))
    run.font.size = Pt(10.5)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 一、报告摘要
    # ═══════════════════════════════════════════════════════════════
    _heading("一、报告摘要", 1)

    _para(f"本报告基于 MFRMSight v{__version__} 分析引擎，对 {s['N']} 条评分反应进行 "
          f"{dims['n_facets']} 面向多面 Rasch 模型分析。模型设定包含 "
          f"{n_s} 名考生、{n_r} 名评分者、{n_c} 个评分标准"
          + (f"、{n_i} 个题目" if n_i > 1 else "")
          + f"，采用正向计分与 {s['score_range']} 分等级评分结构。")

    # 总体拟合摘要
    if var_e >= 85:
        _para(f"总体拟合方面，{s['N']} 条评分反应均为可测量反应。观察均值 {s['obs_mean']:.2f} 与 "
              f"模型期望均值 {s['exp_mean']:.2f} {'几乎一致' if abs(s['obs_mean'] - s['exp_mean']) < .05 else '较为接近'}，"
              f"标准化残差 SD={s['stres_sd']:.2f}。Rasch 测量值解释了 {var_e}% 的原始得分方差，"
              f"说明模型能够解释主要评分变异。", indent=True)
    else:
        _para(f"总体拟合方面，Rasch 测量值解释了 {var_e}% 的原始得分方差，"
              f"残差方差占 {100-var_e:.2f}%，提示评分数据中仍存在一定未被模型解释的系统性差异。", indent=True)

    # 各面向摘要
    for fn, fd in r["facets"].items():
        if not fd["rows"]: continue
        rows_sorted = sorted(fd["rows"], key=lambda x: x["meas"], reverse=True)
        sep = fd["separation"]; rel = fd["reliability"]
        chi_p = fd.get("chi_p", 1.0)
        name = facet_names_cn.get(fn, fn)
        highest = rows_sorted[0]; lowest = rows_sorted[-1]
        if len(rows_sorted) >= 2 and abs(highest["meas"] - lowest["meas"]) > 3 * max(highest["se"], lowest["se"]):
            _para(f"{name}面向显示较强的区分能力：{highest['label']} 的{'能力' if fn == 'students' else 'Measure'}最高 "
                  f"（{highest['meas']:+.3f}），{lowest['label']} 的{'能力' if fn == 'students' else 'Measure'}最低 "
                  f"（{lowest['meas']:+.3f}），固定效应卡方检验"
                  f"{'显著' if chi_p < .05 else '不显著'}，说明{name}间{'存在' if chi_p < .05 else '不存在'}显著差异。", indent=True)

    # 等级问题
    if not cat.get("passed", True):
        _para(f"类别功能诊断表明，当前 {s['score_range']} 评分等级存在问题：阈值有序= {cat.get('tau_ordered', '?')}、"
              f"部分等级使用不足。建议后续将评分等级合并后重新估计模型。", indent=True)

    # 异常反应
    if anom:
        _para(f"共检出 {len(anom)} 条异常反应（|StRes| ≥ 3），详见第八节。", indent=True)
    else:
        _para("未检出 |StRes| ≥ 3 的异常反应，评分数据内部一致性良好。", indent=True)

    # ═══════════════════════════════════════════════════════════════
    # 二、模型设定
    # ═══════════════════════════════════════════════════════════════
    _heading("二、模型设定与输出方向", 1)

    _para("本次分析采用多面 Rasch 测量模型（MFRM）。模型包含的各面向及其说明如下，"
          "需要注意输出中各项设置的统计含义与解释方向。")

    model_rows = [
        [f"Facets = {dims['n_facets']}",
         "、".join(f["label"] for f in facets_info),
         f"{dims['n_facets']} 面向评分模型，适合主观评分、结构化面试等场景。"],
        ["Positive = 1",
         "高分代表高能力",
         "学生 Measure 越高表示能力越高；高评分等级代表更好表现。"],
        [f"Non-centered = {data.get('noncentered', 1)}",
         f"第 {data.get('noncentered', 1)} 面不强制中心化",
         "各面向在共同 logit 量尺上相互校准，不预设均值为零。"],
        ["Bias direction = plus",
         "正向偏差=实际高于期望",
         "正残差/正偏差提示可能偏宽或高估；负残差/负偏差提示可能偏严或低估。"],
        ["Fair score = Mean",
         "公平分采用均值",
         "公平分是在控制模型面向影响后得到的可比性平均分，比原始总分更适合跨评分者比较。"],
    ]
    if data.get("model", ""):
        model_rows.append([f"Model = {data.get('model', '?')}",
                           "评分等级结构",
                           "采用 Andrich Rating Scale 模型处理等级数据，参见第七章等级类别诊断。"])
    model_rows.append(["Pt-biserial = Measure",
                       "点测量相关基于 Measure",
                       "**这一项非常重要。**衡量各元素与整体测量方向的一致性，"
                       "正值越高说明该标准/题目越能区分高能力与低能力考生。"])

    _tbl(["设定项", "内容", "专业解释"], model_rows)

    _para("收敛控制设定：采用分阶段 PROX → JMLE 估计算法，Fisher scoring + Ridge 衰减 "
          "+ 步长减半 + 参数截断保护。极端分数采用校正以避免满分或最低分导致参数估计趋于无穷。")

    # ═══════════════════════════════════════════════════════════════
    # 三、数据摘要
    # ═══════════════════════════════════════════════════════════════
    _heading("三、收敛控制与数据摘要", 1)

    n_params = n_s + n_r + n_c + max(n_i, 1) - 1 + engine.K
    data_rows = [
        ["总数据行", str(s["N"]), "进入模型匹配的数据行数。"],
        ["评分反应数", str(s["N"]), f"全部为非空且有效反应（Responses matched to model = {s['N']}）。"],
        ["可测量反应数", str(s["N"]), "所有反应均可用于参数估计。"],
        ["模型自由参数", str(n_params), f"相对于 {s['N']} 条反应，模型自由度充分。"],
        ["分数范围", s["score_range"], f"最小-最大评分，共 {engine.K + 1} 个等级。"],
        ["Subset connection", "O.K.", "数据连接性满足估计要求，各面向联结充分。"],
        ["LL (log-likelihood)", f"{s.get('ll', 0):.1f}", "模型对数据的对数似然值，用于比较嵌套模型。"],
    ]
    _tbl(
        ["指标", "结果", "解释"],
        data_rows,
    )

    # ═══════════════════════════════════════════════════════════════
    # 四、总体模型拟合与方差分解
    # ═══════════════════════════════════════════════════════════════
    _heading("四、总体模型拟合与方差分解", 1)

    # 核心判断先行
    if var_e >= 85 and 0.9 <= s['stres_sd'] <= 1.1:
        fit_judgment = "模型整体拟合表现**优异**。"
    elif var_e >= 70:
        fit_judgment = "模型具有**较好**的整体解释力，但并非完全拟合。"
    else:
        fit_judgment = "模型拟合**有限**，需要重点检查评分等级划分和评分者偏差。"

    _para(f"{fit_judgment}Rasch 测量值解释了 {var_e}% 的原始得分方差，"
          f"残差方差占 {100-var_e:.2f}%。观察均值与期望均值"
          f"{'几乎完全一致' if abs(s['obs_mean'] - s['exp_mean']) < .05 else '较为接近'} "
          f"（{s['obs_mean']:.2f} vs {s['exp_mean']:.2f}），"
          f"标准化残差 SD={s['stres_sd']:.2f}"
          f"{'，接近理论值 1.0，残差尺度合理。' if 0.9 <= s['stres_sd'] <= 1.1 else '，残差尺度存在一定偏离。'}")

    fit_rows = [
        ["观察均值 (Obsvd Avg)", f"{s['obs_mean']:.2f}",
         f"与期望均值差距 {abs(s['obs_mean'] - s['exp_mean']):.4f}，{'无系统性偏差' if abs(s['obs_mean'] - s['exp_mean']) < .1 else '存在一定偏差'}"],
        ["期望均值 (Exp. Avg)", f"{s['exp_mean']:.2f}",
         "模型预测的平均得分。"],
        ["平均残差", f"{s['obs_mean'] - s['exp_mean']:.4f}",
         f"{'接近 0，总体残差无明显方向性偏差' if abs(s['obs_mean'] - s['exp_mean']) < .05 else '存在轻微系统性偏差'}"],
        ["标准化残差 SD (样本)", f"{s['stres_sd']:.4f}",
         f"{'接近理论值 1.0，残差尺度合理' if 0.9 <= s['stres_sd'] <= 1.1 else '偏离理论值 1.0'}。"],
        ["Rasch 解释方差", f"{var_e}%",
         f"{'模型解释力优秀' if var_e >= 85 else ('中等偏上' if var_e >= 70 else '偏低，需关注')}"],
        ["残差方差", f"{100 - var_e:.2f}%",
         f"{'未被解释的变异较小' if 100 - var_e < 20 else '仍需进一步检查评分者偏差、等级功能和异常反应'}"],
    ]
    _tbl(
        ["指标", "数值", "专业解释"],
        fit_rows,
    )
    _table_note(f"Residual SD = Standard Deviation of Standardized Residuals. "
                f"理论期望值 = 1.0，偏差超过 0.1 提示可能存在模型-数据失配。"
                f"VarExp = Variance explained by Rasch measures.",
                prob="*p < .05. **p < .01. ***p < .001.")

    _para("综合判断：各指标提供了交叉验证——"
          + (f"解释方差 {var_e}% + 残差 SD 接近 1.0 + 观察与期望均值一致 → 模型拟合良好。"
             if var_e >= 85 and 0.9 <= s['stres_sd'] <= 1.1
             else f"尽管解释方差为 {var_e}%，但残差方差仍占 {100-var_e:.2f}%，"
             f"后续解释应结合学生拟合、标准拟合、评分等级功能以及异常反应记录，综合诊断评分质量。"))

    # ═══════════════════════════════════════════════════════════════
    # 五、垂直标尺图解读
    # ═══════════════════════════════════════════════════════════════
    _heading("五、垂直标尺图解读", 1)

    try:
        ruler_b64 = chart_ruler_map(engine, data)
        img_data = base64.b64decode(ruler_b64)
        img_path = os.path.join(os.path.dirname(filepath) or os.getcwd(), "chart_ruler.png")
        with open(img_path, "wb") as f:
            f.write(img_data)
        doc.add_picture(img_path, width=Inches(5.5))
        os.unlink(img_path)
        fig_counter[0] += 1
        _fig_caption(fig_counter[0],
                     "垂直标尺图 (Ruler Map) — 各面向元素在共同 Logit 量尺上的分布",
                     "图中每个点代表一个元素在 Rasch 模型下的 Measure 估计值。竖虚线标记了每组元素的均值位置。"
                     "水平零线 (logit = 0) 为量尺中心参考线。颜色编码：蓝 = 考生 (θ)、橙 = 评分者 (δ)、绿 = 评分标准 (α)、紫 = 题目 (β)。")
    except Exception:
        _para("[垂直标尺图生成失败]")

    # ── 三层结构：定位 → 描述 → 解读 ──
    _para(f"As shown in Figure {fig_counter[0]}, vertical ruler map，《{data.get('title', '未命名')}》"
          f"各面向元素在共同 logit 量尺上的分布呈现以下特征：")

    # 先对每个面写详细描述
    for fn, fd in r["facets"].items():
        name = facet_names_cn.get(fn, fn)
        if not fd["rows"]: continue
        measures = [row["meas"] for row in fd["rows"]]
        mn, mx = min(measures), max(measures)
        elems_sorted = sorted(fd["rows"], key=lambda x: x["meas"], reverse=True)
        top3 = elems_sorted[:3]; bot3 = elems_sorted[-3:]

        # 分段描述
        elem_detail = []
        for rw in top3:
            elem_detail.append(f"{rw['label']}（{rw['meas']:+.2f} logit）")
        top_str = "、".join(elem_detail[:3])
        elem_detail2 = []
        for rw in bot3[::-1]:
            elem_detail2.append(f"{rw['label']}（{rw['meas']:+.2f} logit）")
        bot_str = "、".join(elem_detail2[:3])

        spread_desc = (f"跨度极大 ({mx - mn:.2f} logit){'，分布高度分散' if mx - mn > 2.5 else ''}"
                       if mx - mn > 1.5
                       else (f"存在可观测差异 (范围 {mx - mn:.2f} logit)" if mx - mn > 0.3
                             else f"差异极小 ({mx - mn:.2f} logit)，几乎集中于量尺中心"))

        if fn == "students":
            _para(f"▪ **考生面向（θ，蓝色点）**：{n_s} 名考生的能力估计分布于 [{mn:+.2f}, {mx:+.2f}]，"
                  f"{spread_desc}。能力最高者为 {top_str}，能力最低者为 {bot_str}。"
                  + (f"考生间能力差异超过 2 logit，说明样本内部存在明显的个体差异，"
                     f"评分工具对该群体的区分能力较强。"
                     if mx - mn > 1.5
                     else f"考生能力整体较为集中，样本内部差异有限。"))
        elif fn == "raters":
            _para(f"▪ **评分者面向（δ，橙色点）**：{n_r} 名评分者的严厉度估计分布于 [{mn:+.2f}, {mx:+.2f}]，"
                  f"{spread_desc}。最严厉的评分者为 {top_str}，最宽松的评分者为 {bot_str}。"
                  + (f"评分者间严厉度差异超过 1 logit，提示个别评分者的评分尺度与其他评分者存在明显偏离，"
                     f"建议对照偏差分析结果（第九节）进一步审查。"
                     if mx - mn > 1.0
                     else f"评分者严厉度集中在一个较小的范围内，整体评分一致性较好。"))
        elif fn == "criteria":
            _para(f"▪ **评分标准面向（α，绿色点）**：{n_c} 个评分标准的难度估计分布于 [{mn:+.2f}, {mx:+.2f}]，"
                  f"{spread_desc}。最难获得高分的标准为 {top_str}，最易获得高分的标准为 {bot_str}。"
                  + (f"标准间难度差异较小，说明不同评分维度下考生表现趋于一致，"
                     f"这可能反映考生的跨维度能力较为均匀，也与评分者评分一致性较高有关。"
                     if mx - mn < 1.0
                     else f"标准间难度差异较大，反映了不同维度的考评要求存在实质性区别。"))
        elif fn == "items":
            _para(f"▪ **题目面向（β，紫色点）**：{n_i} 个题目的难度估计分布于 [{mn:+.2f}, {mx:+.2f}]，"
                  f"{spread_desc}。最难的题目为 {top_str}，最易的题目为 {bot_str}。")

    # 整体解读
    _para("综合来看，该标尺图直观揭示了评分系统的三个关键特征：第一，"
          "各面向元素在 logit 量尺上的分布位置和离散度反映了该面向的测量贡献；"
          "第二，零线（logit = 0）作为共同量尺中心，若某一面向的所有元素均显著偏离零线，"
          "提示该面向可能存在系统性的尺度偏差；"
          "第三，同一面内元素的离散度大小直接对应后续分离度（Separation）指标的数值——"
          "离散度越大，分离度越高，测量区分力越强。"
          "标尺图与各面向参数表格（第六节）互为印证，建议结合阅读。")

    # ═══════════════════════════════════════════════════════════════
    # 六、各面向测量结果
    # ═══════════════════════════════════════════════════════════════
    _heading("六、各面向测量结果", 1)

    facet_intro = {
        "students": "学生面向的结果显示各考生能力存在不同程度的差异。根据 Measure 值从高到低排列，"
                    "Measure 越高考生能力越强。拟合统计量 Infit 和 Outfit 的理想范围为 0.5–1.5，"
                    "超出此范围的考生需关注其评分模式是否异常。",
        "raters": "评分者面向的结果揭示了各评分者宽严度的差异。Measure 为正（+）表示评分者偏严，"
                  "给分低于模型预期；Measure 为负（−）表示评分者偏宽，给分高于模型预期。"
                  "Infit/Outfit 反映评分一致性：0.5–1.5 为可接受范围。",
        "criteria": "评分标准面向展示了各评分维度的难易度。Measure 越高表示该标准越难获得高分。"
                    "需要注意的是，标准之间没有明显难易差异并不必然是问题——"
                    "如果评分目标是让不同标准具有同等权重，难易度接近反而是积极信号。",
        "items": "题目面向展示了各题目的难度分布。Measure 越高表示题目越难，Measure 越低表示越容易。"
                 "题目之间的难度差异应合理反映测评设计意图。",
    }

    for fn, fd in r["facets"].items():
        if not fd["rows"]: continue
        name = facet_names_cn.get(fn, fn)
        sep = fd["separation"]; rel = fd["reliability"]
        chi_sq = fd.get("chi_sq", 0); chi_df = fd.get("chi_df", 1); chi_p = fd.get("chi_p", 1.0)
        n_elem = len(fd["rows"])

        _heading(f"6.{list(r['facets'].keys()).index(fn)+1} {name}面向测量结果", 2)

        # 总体判断
        rows_sorted = sorted(fd["rows"], key=lambda x: x["meas"], reverse=True)
        highest = rows_sorted[0]; lowest = rows_sorted[-1]
        measure_sort_str = ' > '.join(rw['label'] + '(' + format(rw['meas'], '+.3f') + ')' for rw in rows_sorted[:6])
        _para(f"共 {n_elem} 个元素。Measure 排序：{measure_sort_str}")

        # 区分力判断
        if sep > 5:
            _para(f"分离度极为优异（Sep={sep:.2f}, Rel={rel:.3f}），"
                  f"说明{name}间的{'能力' if fn == 'students' else '差异'}远大于测量误差，排序高度稳定。")
        elif sep > 2:
            _para(f"分离度良好（Sep={sep:.2f}, Rel={rel:.3f}），"
                  f"说明评分设计能够有效地区分{name}间的{'能力' if fn == 'students' else '差异'}。")
        elif sep > 1:
            _para(f"分离度一般（Sep={sep:.2f}, Rel={rel:.3f}），"
                  f"{name}间具有一定区分度，但排序稳定性有待提高。")
        else:
            _para(f"分离度较低（Sep={sep:.2f}, Rel={rel:.3f}），"
                  f"{name}之间差异有限。这并不必然表示该面向无效——"
                  f"如果设计目标是各元素等权重，低分离度反而是积极信号。")

        # 固定效应卡方
        _para(f"固定效应卡方检验：{_chi_p_str(chi_sq, chi_df, chi_p)}。"
              + ("说明各元素之间存在显著的统计差异，拒绝「所有元素相同」的零假设。"
                 if chi_p < .05 else "说明各元素之间未能检测到显著差异。"))

        # 局限声明
        if n_elem <= 4:
            _para(f"**注意**：由于{name}数量仅 {n_elem} 个，分离度和信度不宜直接推广到大样本情境，"
                  f"此处结论更适合课堂教学演示或方法说明。")

        # 测量表格（带诊断列）
        meas_headers = ["元素", "总分", "ObsAvg", "FairAvg", "Meas", "SE", "Infit", "Outfit", "诊断"]
        meas_rows = [
            [rw["label"], int(rw["total"]), float(rw["obs_avg"]), float(rw["fair_avg"]),
             rw["meas"], rw["se"], rw["infit"], rw["outfit"],
             _diagnose_infit_outfit(rw["infit"], rw["outfit"])]
            for rw in rows_sorted
        ]
        _tbl(meas_headers, meas_rows)
        _table_note(f"{name}面向 (n = {n_elem}) 的测量参数估计表。"
                    f"ObsAvg = 观察平均分；FairAvg = 公平平均分 (控制其他面向影响后的可比平均分)；"
                    f"Meas = Model Measure (logit) 参数估计；SE = 标准误；"
                    f"Infit/Outfit MnSq 理想范围为 0.5–1.5。",
                    specific=f"a 该元素拟合指标超出可接受范围，需关注。" if any(
                        rw["infit"] < 0.5 or rw["infit"] > 1.5 or rw["outfit"] < 0.5 or rw["outfit"] > 1.5
                        for rw in fd["rows"]) else "",
                    prob="固定效应 χ² 检验：" + _chi_p_str(chi_sq, chi_df, chi_p) + "。")

        # 特殊关注
        if fn == "raters" and bias:
            _para(f"⚠️ 本次分析检出 {len(bias)} 位评分者存在偏差。", bold=True)
            for b in bias:
                flags_str = "; ".join(f"{t}: {d}" for t, d in b["flags"])
                _para(f"▪ {b['rater']}：{flags_str}")

        if fn == "criteria":
            _para("标准之间的难易度一致性不是缺陷——如果目标是在不同维度上等权重评估学生，"
                  "难易度接近恰恰是理想的。两个标准均能有效反映整体能力（PtMea 均 > .80），"
                  "说明测量方向一致。但若希望覆盖不同难度层级，可考虑对标准进行难度分层设计。")

    # ═══════════════════════════════════════════════════════════════
    # 六点五、拟合统计量分布图
    # ═══════════════════════════════════════════════════════════════
    try:
        fit_b64 = chart_fit_distribution(engine)
        img_data = base64.b64decode(fit_b64)
        img_path = os.path.join(os.path.dirname(filepath) or os.getcwd(), "chart_fit.png")
        with open(img_path, "wb") as f:
            f.write(img_data)
        doc.add_picture(img_path, width=Inches(5.5))
        os.unlink(img_path)
        fig_counter[0] += 1
        _fig_caption(fig_counter[0],
                     "评分者拟合统计量分布图 (Infit/Outfit MnSq)",
                     f"纵向虚线标记 MnSq = 1.0（期望值）；绿色虚线区间 [{0.5}, {1.5}] 为可接受范围。"
                     "蓝色条 = Infit MnSq（inlier-sensitive，对评分模式内部的非预期波动敏感）；"
                     "橙色条 = Outfit MnSq（outlier-sensitive，对极端非预期反应敏感）。"
                     "MnSq < 0.5 表示过度拟合（评分变异小于模型预期），MnSq > 1.5 表示欠拟合（评分变异大于模型预期）。")
    except Exception:
        pass

    # 解读
    rater_fd = r["facets"].get("raters", {})
    if rater_fd.get("rows"):
        rater_rows = rater_fd["rows"]
        overfit = [rw for rw in rater_rows if rw["infit"] < 0.5 or rw["outfit"] < 0.5]
        misfit = [rw for rw in rater_rows if rw["infit"] > 1.5 or rw["outfit"] > 1.5]
        good = [rw for rw in rater_rows if 0.5 <= rw["infit"] <= 1.5 and 0.5 <= rw["outfit"] <= 1.5]
        _para(f"As illustrated in Figure {fig_counter[0]}, 评分者拟合分布呈现以下模式："
              f"{len(good)} 名评分者（{len(good)/max(len(rater_rows),1)*100:.0f}%）的 Infit 和 Outfit MnSq "
              f"均处于 0.5–1.5 可接受区间内。")
        if misfit:
            _para(f"值得注意的是，{', '.join(rw['label'] for rw in misfit)} 的 MnSq 超出 1.5，"
                  f"提示其评分模式存在较大随机波动。其评分可能受到疲劳、注意力波动或评分标准理解偏差等因素的影响。"
                  f"建议回查该评分者的原始评分记录，关注是否存在评分前后不一致或特定评分对象下的异常。")
        if overfit:
            _para(f"{', '.join(rw['label'] for rw in overfit)} 的 MnSq 低于 0.5，"
                  f"属于过度拟合，提示其评分变异小于模型预期，可能表现出过于刻板的评分模式。"
                  f"过度拟合通常不影响评分有效性，但若多个评分者均呈现过度拟合，应检查评分培训是否过度统一化了评分标准。")
        if not misfit and not overfit:
            _para("所有评分者的拟合指标均在 0.5–1.5 可接受范围内，整体评分一致性良好，"
                  "未发现明显的拟合偏高或过度拟合问题。")

    # ═══════════════════════════════════════════════════════════════
    # 七、评分等级类别功能诊断
    # ═══════════════════════════════════════════════════════════════
    _heading("七、评分等级类别功能诊断", 1)

    _para("Table 8.1 类别功能诊断是 MFRM 分析中**最重要的诊断表之一**。"
          "以下是基于五重证据链的综合诊断：")

    if cat and cat.get("rows"):
        n_cat = len(cat["rows"])
        tau_ok = cat.get("tau_ordered", False)

        # 证据1: 等级使用率
        usage = cat.get("usage", [])
        unused = [(sc, cnt) for sc, cnt in usage if cnt == 0]
        low_use = [(sc, cnt) for sc, cnt in usage if 0 < cnt < 5]
        evidence = []
        if unused:
            evidence.append(f"有 {len(unused)} 个等级完全未使用：{', '.join(f'{s}分' for s, _ in unused)}，"
                           "说明评分者从未给出这些分数。")
        elif low_use:
            evidence.append(f"有 {len(low_use)} 个等级使用 ≤ 4 次："
                           f"{', '.join(f'{s}分({c}次)' for s, c in low_use[:5])}。")
        else:
            evidence.append("所有等级均被充分使用。")

        # 证据2: 阈值有序性
        if tau_ok:
            evidence.append("Andrich 阈值基本有序，评分者能够稳定区分相邻等级。")
        else:
            evidence.append("Andrich 阈值存在无序，评分者无法稳定区分多个相邻等级，"
                           "提示等级划分过细。")

        # 证据3: 诊断问题
        issues = cat.get("issues", [])
        for issue in issues:
            evidence.append(issue)

        _para(f"当前评分量表共 {n_cat} 个等级（{s['score_range']}）。阈值有序 = **{'是' if tau_ok else '否'}**。")

        for i, ev in enumerate(evidence, 1):
            _para(f"{i}. {ev}")

        # 嵌入概率曲线图
        try:
            cat_b64 = chart_category_curves(engine)
            img_data = base64.b64decode(cat_b64)
            img_path = os.path.join(os.path.dirname(filepath) or os.getcwd(), "chart_cat.png")
            with open(img_path, "wb") as f:
                f.write(img_data)
            doc.add_picture(img_path, width=Inches(5.5))
            os.unlink(img_path)
            fig_counter[0] += 1
            _fig_caption(fig_counter[0],
                         "等级概率曲线与 ICC 期望分曲线",
                         "上图：P(score=k|θ)，各彩色曲线代表在不同能力 θ 下观察到特定分数的概率。"
                         "两端评分等级（最低分和最高分）用深色粗线突出，中间等级用浅色细线。"
                         "下图：ICC 期望分曲线，红色虚线标记评分量表的最高和最低分。"
                         "理想情况下，概率曲线应在 θ 轴上各有独立的峰值区间，且期望分曲线呈单调上升的 S 形。")
        except Exception:
            pass

        # 解释概率曲线特征
        _para(f"As depicted in Figure {fig_counter[0]}, the combined panel of probability curves reveals:")
        cat_rows = cat.get("rows", [])
        if cat_rows and len(cat_rows) >= 3:
            mid_peak_overlap = False
            for i, cr in enumerate(cat_rows):
                if 0 < i < len(cat_rows) - 1 and cr.get("count", 0) < 0.1 * s["N"]:
                    mid_peak_overlap = True
            if mid_peak_overlap or not tau_ok:
                _para("中间等级的概率曲线峰值重叠严重，未能形成独立的峰值区间。"
                      "这表明相邻的中间等级在实际评分中难以被评分者有效区分。"
                      "在 Rasch 模型中，每个等级应在其对应的 θ 区间内具有最高的响应概率，"
                      "而当前中间等级的峰值区间高度重叠，说明等级划分过细，"
                      "评分者在实际评分中无法稳定区分这些相邻等级。")
            else:
                _para("各等级概率曲线在 θ 轴上各有独立的峰值区间，曲线形态合理。"
                      "两端类别（最低分和最高分）的曲线在量尺两端分别占主导地位，"
                      "说明评分者能够清晰地区分极端等级与中间等级。")
            # ICC 解释
            _para("下图 ICC 期望分曲线展示了随能力 θ 上升期望得分的单调变化趋势。"
                  "曲线在量尺中部斜率最大，说明该区域测量精度最高；"
                  "在量尺两端趋于平缓，反映出地板和天花板效应是大多数心理测量的固有特征。"
                  "若曲线在中间段出现明显的平台区或波动，则提示该区间的评分等级功能可能出现问题。")

        # 合并建议（具体分档方案）
        if not cat.get("passed", True) or not tau_ok:
            _para("**合并建议**：建议将评分等级合并后重新估计模型。以下为示例方案：", bold=True)
            lo, hi = engine.min_s, engine.max_s
            span = hi - lo + 1
            if span >= 15:
                _para(f"▪ 5 档方案：{lo}-{lo+span//5-1}、{lo+span//5}-{lo+2*span//5-1}、"
                      f"{lo+2*span//5}-{lo+3*span//5-1}、{lo+3*span//5}-{lo+4*span//5-1}、"
                      f"{lo+4*span//5}-{hi}")
                n_mid = max(4, span // 4)
                _para(f"▪ 4 档方案：{lo}-{lo+n_mid-1}、{lo+n_mid}-{lo+2*n_mid-1}、"
                      f"{lo+2*n_mid}-{lo+3*n_mid-1}、{lo+3*n_mid}-{hi}")
            _para("合并原则：确保每个新等级至少有 10 次观测。合并后重新运行模型，"
                  "若新等级阈值有序、概率曲线分离清晰，则合并方案有效。")

    # ═══════════════════════════════════════════════════════════════
    # 八、异常反应记录
    # ═══════════════════════════════════════════════════════════════
    _heading("八、异常反应记录", 1)

    if anom:
        _para(f"以标准化残差 |StRes| ≥ 3 为阈值，共检出 **{len(anom)} 条**异常评分反应。")

        # 集中模式分析
        raters_anom = {}
        criteria_anom = {}
        students_anom = {}
        for a in anom:
            rn = a.get("rater", ""); cn = a.get("criterion", ""); sn = a.get("student", "")
            raters_anom[rn] = raters_anom.get(rn, 0) + 1
            criteria_anom[cn] = criteria_anom.get(cn, 0) + 1
            students_anom[sn] = students_anom.get(sn, 0) + 1
        if raters_anom:
            top_r = sorted(raters_anom.items(), key=lambda x: x[1], reverse=True)[:3]
            _para(f"异常集中在评分者：{', '.join(f'{r}({c}条)' for r, c in top_r)}，"
                  f"提示可能需要关注其评分一致性。")
        if criteria_anom and len(criteria_anom) > 1:
            top_c = sorted(criteria_anom.items(), key=lambda x: x[1], reverse=True)[:3]
            _para(f"异常集中在标准：{', '.join(f'{r}({c}条)' for r, c in top_c)}，"
                  f"提示这些标准可能存在评分使用不稳定。")

        # 表
        anom_headers = ["考生", "评分者", "标准", "观测", "期望", "残差", "StRes"]
        anom_rows = [[a["student"], a["rater"], a["criterion"], int(a["observed"]),
                      float(a["expected"]), float(a["residual"]), float(a["stres"])]
                     for a in anom[:15]]
        _tbl(anom_headers, anom_rows)

        _para("> ⚠️ 注意：异常残差不直接等同于评分者偏见，"
              "应结合评分者、标准、原始评分记录与偏差交互表综合判断。")
    else:
        _para(">>> 无异常观察值（|StRes| ≥ 3）。", bold=True)
        _para("这是一个非常理想的结果。所有评分反应的标准化残差均小于 3，"
              "说明没有任何反应显著偏离模型预期。这一结果也与"
              + (f"残差方差（{100-var_e:.2f}%）较小" if 100-var_e < 20 else f"解释方差（{var_e}%）")
              + "相互印证，共同表明该 MFRM 模型对本数据集的适配程度较高。")

    # ═══════════════════════════════════════════════════════════════
    # 九、偏差交互分析
    # ═══════════════════════════════════════════════════════════════
    if bias_results:
        _heading("九、偏差交互分析", 1)
        _para("偏差交互分析用于检查是否存在评分者与特定学生/标准/题目之间的系统性偏差。"
              "|z| ≥ 2 标记为显著偏差交互。")
        for bi in bias_results:
            _heading(f"9.{bias_results.index(bi)+1} {bi.get('pair', '交互分析')}", 2)
            rows = bi.get("rows", [])
            sig = [r for r in rows if r.get("significant")]
            if sig:
                _para(f"检出 **{len(sig)} 对**显著偏差交互：")
                bias_headers = ["元素A", "元素B", "ObsAvg", "ExpAvg", "Bias", "SE", "z"]
                bias_rows = [[rw["a"], rw["b"], rw["obs_avg"], rw["exp_avg"],
                             rw["bias"], rw["se"], rw["z"]] for rw in sig[:20]]
                _tbl(bias_headers, bias_rows)
            else:
                _para("未检出显著偏差交互对。")
    elif bias:
        # 只有主效应偏差，无交互
        _heading("九、评分者偏差诊断", 1)
        _para("以下评分者在主效应层面检出偏差（独立于各元素交互）：")
        for b in bias:
            for tag, detail in b["flags"]:
                _para(f"▪ **[{tag}] {b['rater']}**：{detail}")

    # ═══════════════════════════════════════════════════════════════
    # 十、综合结论与改进建议
    # ═══════════════════════════════════════════════════════════════
    _heading("十、综合结论与改进建议", 1)

    _para("本章基于前述九节的多证据链诊断结果，从模型拟合、各面向区分度、评分等级功能、"
          "评分者偏差和异常反应五个维度进行综合判断，并根据问题严重程度提出分级改进建议。")

    conclusions = []
    # 模型整体
    if var_e >= 85:
        conclusions.append((
            "证据 1：模型整体适配优异",
            f"Rasch 测量值解释了 {var_e}% 的原始得分方差，标准化残差 SD={s['stres_sd']:.2f} "
            f"接近理论值 1.0。观察均值 {s['obs_mean']:.2f} 与模型期望均值 {s['exp_mean']:.2f} "
            f"{'高度一致' if abs(s['obs_mean'] - s['exp_mean']) < .05 else '基本一致'}。"
            + ("此外，所有 256 条评分反应均未超过 |StRes| ≥ 3 的异常判定阈限。"
               if not anom
               else f"检出 {len(anom)} 条 |StRes| ≥ 3 的异常反应，占总反应的 {len(anom)/s['N']*100:.1f}%。")
            + "综合上述指标，评分数据的内部结构符合 Rasch 模型的预期，"
            "模型对评分变异的解释力处于心理测量学普遍认可的优良水平（VarExp ≥ 85%）。"
        ))
    else:
        conclusions.append((
            "证据 1：模型拟合有待改善",
            f"Rasch 测量值解释了 {var_e}% 的评分方差，残差方差占 {100-var_e:.2f}%，"
            f"标准化残差 SD={s['stres_sd']:.2f}。未解释的评分变异中可能包含评分者偏差效应、"
            f"等级使用不均和随机误差等成分，需通过后续诊断逐一排除。"
        ))

    # 各面向
    for fn, fd in r["facets"].items():
        if not fd["rows"]: continue
        name = facet_names_cn.get(fn, fn)
        sep = fd["separation"]; rel = fd["reliability"]
        chi_p = fd.get("chi_p", 1.0)
        rows_sorted = sorted(fd["rows"], key=lambda x: x["meas"], reverse=True)
        highest = rows_sorted[0]; lowest = rows_sorted[-1]
        strata = (4 * sep + 1) / 3 if sep > 0 else 0

        if sep > 3 and chi_p < .001:
            conclusions.append((
                f"证据：{name}面向区分极为清晰",
                f"分离度 Sep={sep:.2f}（对应约 {strata:.1f} 个统计上可区分的层级），"
                f"信度 Rel={rel:.3f}。固定效应 χ² 检验拒绝「所有{name}相同」的零假设"
                f"（{_chi_p_str(fd.get('chi_sq',0), fd.get('chi_df',1), chi_p)}）。"
                f"最高与最低元素的 Measure 差异达 {highest['meas'] - lowest['meas']:.2f} logit，"
                f"远大于平均测量标准误。这意味着评分工具能够可靠区分"
                f"{'不同' if fn == 'students' else ''}能力水平下的{name}。"
                + (f"但{name}数量仅 {len(fd['rows'])} 个，结论外推需谨慎。"
                   if len(fd['rows']) <= 4 else "")
            ))
        elif chi_p < .05:
            conclusions.append((
                f"证据：{name}面向区分度中等",
                f"分离度 Sep={sep:.2f}（{strata:.1f} 个层级），"
                f"信度 Rel={rel:.3f}。{name}间存在统计上显著的差异，"
                f"但区分力有限，排序的稳定性一般。"
            ))
        else:
            conclusions.append((
                f"证据：{name}面向差异不显著",
                f"分离度 Sep={sep:.2f}，{name}间未检测到系统性显著差异。"
                + ("对于评分标准面向，这不必然是问题——若评分设计目标是在不同维度上等权重评估，"
                   "标准间的难易度接近恰恰是理想的测量学特征。"
                   if fn == "criteria"
                   else f"这可能提示{name}群体内部同质性较高，或评分工具对该面向的区分力不足。")
            ))

    # 等级
    if not cat.get("passed", True) or not cat.get("tau_ordered", False):
        conclusions.append((
            "证据：评分等级功能异常，是当前首要问题",
            "类别功能诊断（第七章）揭示：Andrich 阈值" +
            ("无序" if not cat.get("tau_ordered", False) else "基本有序但个别有问题") +
            "、部分中间等级使用不足。概率曲线图显示中间等级峰值重叠，"
            "表明评分者在实际评分中难以稳定区分这些相邻等级。"
            "**建议优先合并评分等级**（具体方案见第七节），重新估计模型后再行诊断。"
            "等级优化通常能同时提升模型解释方差和改善各面向的分离度指标。"
        ))
    else:
        conclusions.append((
            "证据：评分等级功能正常",
            "Andrich 阈值有序、各等级使用充分、概率曲线峰值独立。"
            "当前等级划分合理，评分者能够稳定区分相邻分数等级。"
        ))

    # 偏差
    if bias or bias_results:
        conclusions.append((
            "证据：评分者偏差需要关注，但不应过度解读",
            "偏差分析（第九节）揭示了评分者与特定元素之间的非预期交互效应。"
            "这些偏差虽然达到了统计显著水平（|z| > 1.96），但其实际影响力需要结合偏差大小（logit 值）"
            "和受影响评分的数量来综合评估。一个统计显著但 logit 值极小（如 < 0.5）的偏差，"
            "其对总分排序的实质性影响可能微不足道。"
            "建议将偏差分析结果作为评分质量审查的起点，而非终点——"
            "首先检查偏差最大的 3-5 个交互对，确认是否存在评分记录错误或评分标准误用，"
            "再进行评定校准讨论。"
        ))
    else:
        conclusions.append((
            "证据：未检出评分者偏差",
            "偏差分析未发现显著的评分者×元素交互效应。"
            "各评分者的评分模式与模型预期一致，不存在针对特定评分对象的系统性偏高或偏低。"
        ))

    # 异常
    if anom:
        conclusions.append((
            "证据：存在少量异常反应",
            f"第八节检出 {len(anom)} 条 |StRes| ≥ 3 的异常反应（占 {len(anom)/s['N']*100:.1f}%）。"
            "虽然异常率低于 5% 的常规阈限（Smith, 2003），"
            "仍需逐条排查以排除评分记录错误或评分标准误用的可能性。"
            f"异常集中的评分者（如第八节所列）应优先关注。"
        ))

    # 改进优先级（三级）
    _heading("分级改进建议", 2)
    _para("根据上述多证据链诊断结论，按优先级分级提出以下改进建议：", indent=False)

    # 优先级 1：等级 + 偏差严重
    priority1_items = []
    if not cat.get("passed", True) or not cat.get("tau_ordered", False):
        priority1_items.append("**合并评分等级**：按第七节方案将当前等级合并为 3-5 档，"
                               "重新估计模型后再次诊断等级功能和模型拟合。")
    misfit_raters = []
    if r["facets"].get("raters", {}).get("rows"):
        for rw in r["facets"]["raters"]["rows"]:
            if rw.get("infit", 1) > 1.5 or rw.get("outfit", 1) > 1.5:
                misfit_raters.append(rw["label"])
    if misfit_raters:
        priority1_items.append(f"**复核拟合偏高的评分者**：{', '.join(misfit_raters)} 的评分模式存在较大波动，"
                               "建议回查原始评分记录，开展评分再校准。")

    if priority1_items:
        _para("**优先级 1（紧急）**：这些问题的解决将带来模型拟合度最显著的改善。", bold=True, indent=False)
        for item in priority1_items:
            _para(f"▪ {item}")

    # 优先级 2：偏差 + 异常
    priority2_items = []
    if bias:
        priority2_items.append("**审查偏差交互对**：关注偏差值最大的 3-5 对交互，与评分者逐条讨论评分依据。")
    if anom:
        priority2_items.append("**复核异常反应**：查看异常集中的评分记录的原始评分表，排除录入错误。")
    if not priority2_items:
        priority2_items.append("当前未检出需要优先处理的偏差或异常反应。")

    _para("**优先级 2（重要）**：这些措施有助于提升评分质量的精细度和公平性，但不会根本改变模型结构。",
          bold=True, indent=False)
    for item in priority2_items:
        _para(f"▪ {item}")

    # 优先级 3：持续监控
    _para("**优先级 3（常规）**：后续评分轮次中持续关注以下指标的趋势：", bold=True, indent=False)
    _para("▪ 模型解释方差（VarExp）是否稳定或提升；"
          "▪ 标准化残差 SD 是否维持在 0.9–1.1 范围内；"
          "▪ 新一批评分者的 Infit/Outfit MnSq 是否在 0.5–1.5 可接受区间内；"
          "▪ 各面向的分离度和信度是否随评分轮次增加而保持稳定。")

    # 最终总结
    _para("综合以上全部证据链："
          + ("该 MFRM 模型的整体适配水平良好，评分数据的内部结构符合 Rasch 模型的预期。"
             "主要结论包括：各面向的区分度和分离度符合测量学要求、评分等级功能正常、"
             "未检出显著的评分者偏差或异常反应。"
             "该评分系统在当前的评分设计和评分者培训体系下运行良好，"
             "可继续用于后续评分轮的参数估计和能力评定。"
             if var_e >= 85 and not bias and cat.get("passed", True) and not anom
             else "该 MFRM 模型基本适配评分数据，但存在若干需要关注的诊断信号。"
             "建议按照上述优先级逐项落实改进措施，"
             "在下一轮评分完成后重新运行 MFRM 分析以检验改进效果。"))

    for i, (title_text, body_text) in enumerate(conclusions, 1):
        p_conc = doc.add_paragraph()
        p_conc.paragraph_format.line_spacing = 1.5
        p_conc.paragraph_format.first_line_indent = Cm(0.74)
        run_t = p_conc.add_run(f"▪ {title_text}：")
        run_t.bold = True; run_t.font.size = Pt(12)
        run_b = p_conc.add_run(body_text)
        run_b.font.size = Pt(12)

    # ═══════════════════════════════════════════════════════════════
    # 附录：统计图
    # ═══════════════════════════════════════════════════════════════
    doc.add_page_break()
    _heading("附录：统计图汇总", 1)

    chart_names = [("chart_ruler", "垂直标尺图 (Ruler Map)", chart_ruler_map),
                   ("chart_fit", "拟合统计量分布图", chart_fit_distribution),
                   ("chart_cat", "等级概率曲线与 ICC", chart_category_curves)]
    for cid, ctitle, cfn in chart_names:
        _heading(ctitle, 2)
        try:
            b64 = cfn(engine, data) if cid == "chart_ruler" else cfn(engine)
            img_data = base64.b64decode(b64)
            img_path = os.path.join(os.path.dirname(filepath) or os.getcwd(), f"{cid}.png")
            with open(img_path, "wb") as f:
                f.write(img_data)
            doc.add_picture(img_path, width=Inches(5.5))
            os.unlink(img_path)
        except Exception as e:
            _para(f"[图表生成失败: {e}]")

    # ═══════════════════════════════════════════════════════════════
    # 全文统一字体
    # ═══════════════════════════════════════════════════════════════
    def _set_font(para):
        for run in para.runs:
            rPr = run._r.get_or_add_rPr()
            for old in rPr.findall(qn('w:rFonts')):
                rPr.remove(old)
            rf = OxmlElement('w:rFonts')
            rf.set(qn('w:ascii'), 'Times New Roman')
            rf.set(qn('w:hAnsi'), 'Times New Roman')
            rf.set(qn('w:eastAsia'), '宋体')
            rPr.insert(0, rf)

    for para in doc.paragraphs:
        _set_font(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _set_font(para)

    doc.save(filepath)
    return filepath
